"""生成 AI Agent 工程师实习生面试 — 项目经历文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.size = Pt(10.5)
style.font.name = 'Microsoft YaHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def h1(text):
    doc.add_page_break()
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.size = Pt(16)
    return h

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.size = Pt(13)
    return h

def h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.size = Pt(11)
    return h

def body(text):
    p = doc.add_paragraph(text)
    for r in p.runs: r.font.size = Pt(10.5)
    return p

def bold_body(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True; r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs: r.font.size = Pt(10.5)
    return p

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def highlight(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    return p

# ══════════════════════════════════════════
# 封面
# ══════════════════════════════════════════
for _ in range(3): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('AI Agent 工程师实习生 — 项目经历')
r.font.size = Pt(24); r.bold = True; r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('EAIOS 企业智能操作系统 · 独立开发全栈项目')
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph()
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f'{datetime.date.today().strftime("%Y-%m-%d")}   |   面试准备用').font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════
# Chapter 1: 项目概述
# ══════════════════════════════════════════
h1('一、项目概述')

h2('1.1 一句话介绍')
body('独立设计并实现了一个企业级 AI 操作系统（EAIOS），核心是 LangGraph 驱动的 Multi-Agent 编排引擎，16 个角色化 Agent 在 Supervisor Pattern 下协作，结合 RAG 知识库、数字孪生、预测引擎和主动告警，完成企业从"数据录入"到"智能决策"的完整闭环。')

h2('1.2 技术栈速览')
tech_items = [
    'Agent 框架：LangGraph StateGraph + LangChain + Supervisor Pattern',
    'LLM：DeepSeek API（OpenAI 兼容）/ BGE Embedding（512 维）',
    '后端：FastAPI + SQLAlchemy 2.0 Async + Celery + Redis + pgvector',
    '前端：Next.js 15 + React 19 + TypeScript + TailwindCSS + shadcn/ui',
    '数据：PostgreSQL + pgvector（向量检索）/ Neo4j（知识图谱）',
    '部署：Docker Compose（6 服务：Nginx/Frontend/Backend/Postgres/Redis/Neo4j）',
    '协议：WebSocket（流式对话）+ SSE（决策分析）+ REST（API）',
]
for item in tech_items:
    bullet(item)

h2('1.3 项目规模')
bold_body('代码量：', '~20,000+ 行（后端 ~8,000 / 前端 ~8,000 / 配置 ~4,000）')
bold_body('Agent 数量：', '16 个（7 部门 Agent + 9 通用 Agent）')
bold_body('数据库表：', '22 张（PostgreSQL）+ 7 实体 × 8 关系（Neo4j）')
bold_body('API 端点：', '19 个路由模块，100+ REST 端点')
bold_body('前端页面：', '16 个（Chat / Documents / Knowledge / Graph / Memories / Analytics / Forecast / Decision Center / Reports / Members / 6 Setting 页）')
bold_body('Docker 服务：', '6 个（Postgres / Redis / Neo4j / Backend / Frontend / Nginx）')

doc.add_page_break()

# ══════════════════════════════════════════
# Chapter 2: Multi-Agent 系统
# ══════════════════════════════════════════
h1('二、Multi-Agent 编排系统（核心亮点）')

h2('2.1 架构设计')
body('采用 LangGraph StateGraph + Supervisor Pattern 构建 16 个 Agent 的统一编排系统，支持单 Agent 执行和多 Agent 流水线。')
body('')

bold_body('Supervisor 路由：', 'LLM 分析用户意图 → 输出 JSON 路由决策 → LangGraph 按序调度 Agent。LLM 失败时 Keyword Fallback 兜底。')
bold_body('AgentState（共享状态）：', 'TypedDict 定义：messages / user_query / next_agents / context_text / sources / final_response / agent_trace。Agent 间通过 State 传递上下文，无状态隔离问题。')
bold_body('流水线执行：', 'search → analyst → writer 三段式：检索文档 → 数据分析 → 生成报告，前一个 Agent 的输出自动传入下一个 Agent。')

h2('2.2 16 个 Agent 分类')
bullet('通用 Agent（9 个）：chat / search / research / analyst / writer / sql / profile / memory / decision')
bullet('部门 Agent（7 个）：CEO / CFO / VP Sales / CHRO / COO / CPO / VP Customer')
bullet('每个 Agent = 专属 System Prompt（角色 + 分析框架 + 输出格式）+ run_X_agent(query, context) 函数 + LangGraph Node 函数')

h2('2.3 上下文注入（Agent 感知企业数据的关键）')
body('在 Agent 路由决策之后、Agent 执行之前，Orchestrator 自动注入 6 层上下文：')
bullet('Step 5：RAG 检索（pgvector 余弦搜索 → Top-5 文档 Chunk）')
bullet('Step 5.5：记忆召回（重要性 × 40% + 时效性 × 30% + 向量相似度 × 30%）')
bullet('Step 5.6：企业画像（Company / Department / Employee 等结构化数据）')
bullet('Step 5.7：文档列表（文件名匹配）')
bullet('Step 5.8：数字孪生指标（Revenue / Cost / Customer Count 等时序数据）')
body('这让每个 Agent 获得对企业全貌的感知能力，而非仅依赖用户输入的一句话。')

h2('2.4 流式返回')
body('对话走 WebSocket（token 级流式），决策分析走 SSE。自定义 WebSocketClient 带指数退避自动重连（3s→30s，最多 10 次）+ 30s Ping 心跳。WebSocket 失败时自动回退到 HTTP SSE。')

h2('2.5 对话后处理')
bullet('自动记忆提取：LLM 判断对话中是否有值得记住的事实/决定 → JSON → BGE Embedding → pgvector 去重入库')
bullet('Agent Trace 记录：完整的 Agent 执行链（如 ["search", "finance", "ceo"]）')

doc.add_page_break()

# ══════════════════════════════════════════
# Chapter 3: RAG & Tool 系统
# ══════════════════════════════════════════
h1('三、RAG 检索增强生成 & Tool 系统')

h2('3.1 RAG 流水线')
body('设计了两阶段 RAG 流水线：')
bold_body('离线阶段：', '文档上传 → PyMuPDF/Python-docx/Python-pptx 多格式解析 → tiktoken 512 chunk 切割（50 token overlap）→ BGE-small-zh-v1.5 Embedding → pgvector 存储')
bold_body('在线阶段：', 'Query Rewrite（LLM 改写查询）→ pgvector 余弦相似度搜索 → Top-5 召回 → LLM 合成带引用答案')

h2('3.2 Tool 系统设计')
body('实现了 ToolRegistry 单例模式，统一管理 10 个 LangChain Tool：')
bullet('Document Search Tool：搜索上传文档')
bullet('Web Search Tool：Tavily API 联网搜索')
bullet('SQL Query Tool：自然语言转 SQL 查询数据库')
bullet('Business Metrics Tool：查询数字孪生指标和趋势（实际代码查询 business_metrics 表）')
bullet('Company Profile Tool：查询企业画像数据')
bullet('Memory Tool：查询/存储企业记忆')
bullet('Graph Tool：查询 Neo4j 知识图谱')
bullet('更多 ...')

h2('3.3 知识图谱')
body('Neo4j 存储企业实体和关系（Company / Department / Employee / Product / Customer / Process / Goal 共 7 种实体 + 8 种关系类型）。')
bullet('分词搜索：用户查询 → tokenize（"营销部主管是谁" → ["营销部", "主管"]）→ Cypher OR CONTAINS 查询 → 返回匹配实体')
bullet('安全：Cypher 白名单（只允许 MATCH / RETURN / CALL）')

doc.add_page_break()

# ══════════════════════════════════════════
# Chapter 4: 技术难点
# ══════════════════════════════════════════
h1('四、技术难点与解决方案')

h2('4.1 Supervisor 路由的可靠性')
bold_body('问题：', '纯 LLM 路由可能因 API 超时/JSON 解析失败而导致系统不可用。')
bold_body('方案：', '实现双重路由策略 —— LLM 路由（主） + Keyword Fallback（备）。每个 Agent 定义触发关键词列表，LLM 失败时自动降级为关键词匹配。这种设计让系统在 LLM 故障时仍然可用。')

h2('4.2 Agent 流水线的上下文传递')
bold_body('问题：', '多个 Agent 顺序执行时，如何让后面的 Agent 获得前面 Agent 的发现？')
bold_body('方案：', 'LangGraph StateGraph 的共享 State 机制。每个 Agent Node 函数签名是 (AgentState) → AgentState，前一个 Node 修改 State 后自动传入下一个 Node。agent_trace 记录完整执行链，context_text 在流水线中逐步累加。')

h2('4.3 流式对话的容错')
bold_body('问题：', 'WebSocket 可能因为网络波动断开。')
bold_body('方案：', '自定义 WebSocketClient 实现三层容错 —— 指数退避重连（3s/6s/12s/24s/30s 最大 10 次）+ Ping 心跳检测 + WebSocket 失败时自动回退 HTTP SSE。')

h2('4.4 高并发下的数据库优化')
bold_body('问题：', 'FastAPI 异步 + pgvector 向量检索需要高效的连接管理。')
bold_body('方案：', 'SQLAlchemy 2.0 async engine 配置连接池（pool_size=20 / max_overflow=10）+ AsyncSession 依赖注入（每个请求独立 session）+ pgvector IVFFlat 索引加速近似搜索。')

h2('4.5 环境差异导致的前端 API 路由失效')
bold_body('问题：', '开发环境（前后端不同端口直连）和生产环境（Nginx 统一入口）的 API URL 逻辑不同。')
bold_body('方案：', '前端使用相对路径 API_BASE="" 替代硬编码绝对 URL。axios 发出的请求走同源，由 Nginx 代理到后端。同一份构建产物在任何域名下都能工作。基于项目实际踩坑经验总结。')

h2('4.6 Docker 环境下 Python 依赖治理')
bold_body('问题：', 'sentence-transformers 拉入完整 CUDA 版 PyTorch 导致 Docker 镜像膨胀到 5GB+ 且构建超时。')
bold_body('方案：', 'Dockerfile 中先安装 CPU-only PyTorch 再安装其他依赖，镜像体积缩减 60%+。同时配置国内镜像源（阿里云 Debian + PyPI）解决网络超时问题，passlib + bcrypt 版本兼容性修复。')

doc.add_page_break()

# ══════════════════════════════════════════
# Chapter 5: 我能聊什么
# ══════════════════════════════════════════
h1('五、面试可深入讨论的话题（STAR 法则准备）')

body('以下每个话题对应一个你亲身解决过的技术问题，用 STAR 法则（情境-任务-行动-结果）组织回答。')

h2('5.1 Agent 编排')
bullet('"为什么选择 Supervisor Pattern 而不是完全自主的 Multi-Agent 辩论？"')
bullet('"如果新增一个 Agent，需要改哪些文件？如何设计它的 System Prompt？"')
bullet('"Supervisor 路由失败时系统如何兜底？Keyword Fallback 的具体实现？"')
bullet('"多个 Agent 之间如何共享信息？AgentState 的数据结构为什么这样设计？"')

h2('5.2 RAG & 向量检索')
bullet('"从文档上传到 Agent 能检索到的完整流程是怎样的？"')
bullet('"为什么使用 pgvector 而不是独立的向量数据库（Pinecone/Milvus）？"')
bullet('"如果检索结果不够好，你有哪些优化方向？"')
bullet('"Query Rewrite 的作用是什么？为什么需要这个步骤？"')

h2('5.3 LangGraph & LangChain')
bullet('"LangGraph 的 StateGraph 和 Chain 的区别？什么场景用哪个？"')
bullet('"Tool 是如何注册和发现给 Agent 的？ToolRegistry 的设计？"')
bullet('"Structured Output 在项目中用在哪里？Router 的路由决策如何保证 JSON 格式有效？"')

h2('5.4 工程实践')
bullet('"JWT 认证流程 + Token 自动刷新（并发防抖）是如何实现的？"')
bullet('"Docker Compose 下 6 个服务的启动顺序和健康检查策略？"')
bullet('"Celery + Celery Beat 的定时任务如何配置？为什么用 Redis 做 broker？"')
bullet('"前后端部署中遇到了哪些实际踩坑问题？如何排查和修复的？"')

h2('5.5 系统设计思路')
bullet('"如果让你从头设计一个 Multi-Agent 客服系统，你会怎么设计？"')
bullet('"当前项目的上下文注入 6 层有没有冗余或缺失？"')
bullet('"Agent 的记忆系统（自动提取 + 衰减 + 混合召回）的设计理念？"')

doc.add_page_break()

# ══════════════════════════════════════════
# Chapter 6: 面试话术
# ══════════════════════════════════════════
h1('六、面试话术参考')

h2('6.1 自我介绍（2 分钟版本）')
body('"我最近独立完成了一个全栈项目——企业 AI 操作系统 EAIOS。核心是一个 LangGraph 驱动的 Multi-Agent 编排系统，包含 16 个角色化 Agent，采用 Supervisor Pattern 进行路由和流水线执行。项目覆盖了完整的 AI 工程链路：RAG 检索增强生成（pgvector + BGE Embedding）、知识图谱（Neo4j + Cypher）、数字孪生（时序业务数据自动注入 Agent 上下文）、预测引擎、主动告警和战略决策中心。技术栈上后端用 FastAPI + SQLAlchemy Async，前端 Next.js 15，6 个 Docker 服务编排部署。代码量两万行以上，22 张数据库表，19 个 API 路由模块，100+ REST 端点。整个过程让我深入理解了 LangGraph 的 StateGraph 架构、Agent 上下文注入策略、RAG 流水线的设计权衡，以及大规模 Python 项目的工程化实践。"')

h2('6.2 "为什么选择 LangGraph 而不是自己实现 Agent 编排？"')
body('"LangGraph 提供了两个关键能力：一是 StateGraph 的状态管理——Agent 之间通过共享 TypedDict State 传递上下文，添加新 Agent 只需实现 (State) → State 的 Node 函数并注册到 Graph，代码侵入性极低；二是它内置了条件边（conditional edges）支持 Supervisor 路由决策后的动态调度——LLM 返回路由结果后，LangGraph 自动走对应的边到目标 Agent Node。自己实现这些需要处理状态持久化、执行顺序控制、流式输出中断/恢复等大量边缘情况。用 LangGraph 让我们把精力集中在 Agent 能力和上下文设计上，而不是编排框架本身。"')

h2('6.3 "项目中最大的技术挑战是什么？"')
body('"最大的挑战是让 Agent 感知企业全貌。纯 LLM 只知道用户输入的那句话，但企业级 Agent 需要知道这个企业的收入数据、组织架构、历史决策和文档知识。我设计了一个 6 层上下文注入机制：在 Supervisor 路由之后、Agent 执行之前，Orchestrator 自动从 5 个数据源（RAG 检索、企业记忆、企业画像、文档列表、数字孪生指标）拉取上下文，拼接成一段结构化文本注入到 Agent 的 System Prompt 前面。关键是这些操作是异步并行的，不显著增加响应延迟。另一个挑战是路由可靠性——纯 LLM 路由可能在 API 故障时使整个系统不可用，所以我加了 Keyword Fallback 作为降级方案，让系统在 LLM 不可用时至少能通过关键词匹配找到合适的 Agent。"')

h2('6.4 "你对 AI Agent 的理解是什么？"')
body('"我认为 AI Agent 有三个核心要素：感知（Perceive）、推理（Reason）、行动（Act）。感知不只是理解用户输入，还包括自动获取相关上下文——就像我这个项目里的 6 层上下文注入。推理是 LLM 能力 + System Prompt 引导 + Tool 调用——Agent 需要知道自己能调用哪些工具，以及何时调用。行动不只是生成文本，还可以是数据库查询、文档搜索、调用外部 API、触发告警。在这个项目里，不同的 Agent 有不同粒度的行动能力：search Agent 能检索文档，sql Agent 能查询数据库，finance Agent 能分析财报数据。每个 Agent 是一个专门领域的推理+行动单元，Supervisor 负责把多个 Agent 编排成解决问题的流水线。"')

h2('6.5 "你还有什么想问我们的？"')
bullet('"团队目前在 Agent 编排上用什么框架？LangGraph 还是自研方案？"')
bullet('"Agent 的上下文管理——你们如何处理长对话中的记忆衰减和重要信息提取？"')
bullet('"Agent 的 Test & Eval——你们如何评估一个 Agent 的输出质量？"')
bullet('"模型选型——团队主要用哪些 LLM？有没有做 Fine-tuning 或 RAG 优化？"')

doc.add_page_break()

# ══════════════════════════════════════════
# Chapter 7: 技术词汇速查
# ══════════════════════════════════════════
h1('七、面试高频技术词汇速查表')

body('以下词汇表帮助你在面试中准确使用技术术语，避免口语化或含糊表述。')

vocab = [
    ('Supervisor Pattern', 'Multi-Agent 编排的一种模式：一个 Supervisor Agent 负责分析任务、路由到合适的 Worker Agent、协调执行顺序。区别于完全自主的 Debate Pattern 或层级式的 Hierarchical Pattern。'),
    ('StateGraph', 'LangGraph 的核心抽象：一个有向图，节点是 Agent 函数，边是状态传递路径。每个节点接收 State、返回 State，图引擎负责按序执行。'),
    ('RAG（Retrieval-Augmented Generation）', '检索增强生成：先检索相关文档，再将检索结果作为上下文注入 LLM Prompt，提升回答的事实性和时效性。'),
    ('Embedding', '将文本映射为固定维度的浮点数向量，语义相似的文本向量接近。BGE Embedding 是 BAAI 开源的中英文嵌入模型。'),
    ('pgvector', 'PostgreSQL 的向量扩展，新增 vector 数据类型和余弦距离（<=>）运算符，让关系数据库原生支持向量检索。'),
    ('Cypher', 'Neo4j 的图查询语言，类似于 SQL 但用于查询图结构（节点 + 关系 + 路径）。'),
    ('Celery + Beat', 'Celery 是 Python 的分布式任务队列，Beat 是它的定时调度器。项目用 Redis 做消息代理，实现每小时/每天的定时告警检测。'),
    ('SSE（Server-Sent Events）', '服务器向客户端单向推送事件的 HTTP 协议。适合长文本生成（如决策分析），比 WebSocket 更轻量。'),
    ('JWT（JSON Web Token）', '无状态认证令牌，包含用户信息、过期时间、签名。项目实现 access_token（60min）+ refresh_token（7d）双 Token 机制。'),
    ('Docker Compose', '多容器编排工具，通过 YAML 定义服务依赖、网络、存储卷。项目 6 个服务通过 depends_on + healthcheck 控制启动顺序。'),
    ('TypedDict', 'Python typing 模块的类型提示，定义字典的键和值类型。LangGraph 用它定义 AgentState 的结构。'),
    ('Tool / Function Calling', 'LLM 不直接生成文本，而是输出结构化 JSON 调用预定义的函数（Tool）。LangChain 封装了 Tool 的定义、参数校验和结果解析。'),
]
for term, desc in vocab:
    h2(term)
    body(desc)

doc.add_page_break()

# ══════════════════════════════════════════
# 保存
# ══════════════════════════════════════════
doc.save('EAIOS_Project_Experience.docx')
print('Saved: EAIOS_Project_Experience.docx')