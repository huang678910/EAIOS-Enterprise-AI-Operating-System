"""Generate EAIOS Technical Documentation Word Document"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.size = Pt(10.5)

# Title Page
for _ in range(3): doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('EAIOS — Enterprise AI Operating System')
run.font.size = Pt(26); run.bold = True; run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('企业智能操作系统 — 技术实现文档')
r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'v1.0  |  {datetime.date.today().strftime("%Y-%m-%d")}  |  16 Agent / 19 API / 16 Pages / 22 Tables').font.size = Pt(11)
doc.add_page_break()

# Helper
def h2(text): return doc.add_heading(text, level=2)
def h1(text):
    doc.add_page_break()
    return doc.add_heading(text, level=1)
def bullet(text): doc.add_paragraph(text, style='List Bullet')
def para(text): doc.add_paragraph(text)

# ═══ Ch1 ═══
h1('1. 项目概述')
para('EAIOS是一个企业专属智能操作系统，核心能力：理解企业、记忆企业、分析企业、预测企业、辅助企业决策。')
para('16 Agent | 19 API路由 | 16前端页面 | 22张数据表 | 6 Docker服务')

# ═══ Ch2 ═══
h1('2. 技术栈全景')
h2('2.1 前端')
for t,d in [('Next.js 15','React全栈框架, App Router'),('React 19','UI组件库'),('TypeScript','严格模式'),('TailwindCSS','原子化CSS'),('Zustand','状态管理'),('Recharts','数据图表'),('Axios','HTTP客户端, Token自动刷新'),('WebSocket','实时流式响应')]:
    bullet(f'{t}: {d}')

h2('2.2 后端')
for t,d in [('FastAPI','异步Web框架'),('SQLAlchemy 2.0 Async','异步ORM + pgvector'),('Pydantic v2','数据校验'),('Alembic','数据库迁移(13个)'),('Celery + Beat','异步任务 + 定时调度'),('Redis','缓存 + 消息代理'),('LangGraph','多Agent编排(StateGraph)'),('LangChain','LLM工具链'),('DeepSeek API','主力LLM'),('BGE Embedding','中文向量(512维)'),('pgvector','PostgreSQL向量扩展'),('Neo4j 5','图数据库'),('Whisper','音频转录'),('PyMuPDF','PDF解析')]:
    bullet(f'{t}: {d}')

# ═══ Ch3: 10 Layers ═══
h1('3. 10层架构与功能实现')

layers = [
    ('L1 企业画像', [
        '9张表: companies/departments/positions/employees/products/customers/processes/goals/kpis',
        '37 REST端点, 4前端页面(公司/组织/业务/目标)',
        '关键逻辑: Goals方向判断(higher/lower) + progress_pct自动计算; Departments树形自引用'
    ]),
    ('L2 知识中心', [
        '文档流水线: 上传→解析(10格式)→Chunk(512)→Embedding→pgvector',
        'RAG检索: Query Rewrite(LLM)→pgvector余弦搜索→Reranker→LLM合成',
        '6个连接器: GitHub/Notion/Jira/Confluence/Slack/WeCom, 统一BaseConnector接口',
        '音频转录: Whisper small/base→文本入库→RAG可检索'
    ]),
    ('L3 企业记忆', [
        '三种类型: Long-Term(长期)/Episodic(事件)/Semantic(语义)',
        '自动提取: 对话后LLM判断→JSON→pgvector Embedding→去重→入库',
        '混合加权召回: importance×40% + recency×30% + vector×30%',
        '衰减: 30天未访问→importance -0.5'
    ]),
    ('L4 知识图谱', [
        'Neo4j: 7种实体 + 8种关系, 自动同步Layer1数据',
        '分词搜索: tokenize("谁在营销部主管")→["营销部","主管"]→CONTAINS任一关键词',
        '安全: Cypher白名单(MATCH/RETURN/CALL)'
    ]),
    ('L5 数字孪生', [
        'business_metrics表: metric_name + value + unit + period + category',
        'CRUD API: record/batch/snapshot/trend/summary(6端点)',
        '排序: category→metric_name→recorded_at DESC',
        '上下文注入: Orchestrator Step5.8自动注入Agent context_text'
    ]),
    ('L6 部门Agent', [
        '7个角色化Agent: CEO/CFO/VP Sales/CHRO/COO/CPO/VP Customer',
        '模式: 专属SystemPrompt→run_X_agent(query,context)→LLM streaming→final_response',
        'Supervisor路由: LLM路由(primary)→Keyword Fallback→pipeline顺序执行',
        '上下文: 自动注入企业画像+Memory+Metrics+RAG文档'
    ]),
    ('L7 分析中心', [
        'AnalyticsService: 聚合metrics snapshot+trends+KPIs+Goals+alerts',
        '业务排序: revenue→cost→hr→inventory→customer→operations',
        'AI分析: LLM→Summary+Key Insights+Recommendations→缓存1小时',
        '前端: KPI卡片(Recharts)+趋势图+进度条+告警面板'
    ]),
    ('L8 预测引擎', [
        '两种方法: 移动平均(加权)+线性回归(sklearn LinearRegression)',
        '自动推算: periods=min(历史数/2,6), 清除旧预测去重',
        '置信区间: ±1.5×标准差',
        '前端: 指标去重下拉→方法→Predict→三线图'
    ]),
    ('L9 主动式AI', [
        '三层检测: 规则引擎(KPI阈值)+统计异常(Z-score/环比/反转)+AI评估(LLM过滤)',
        'Celery Beat: 每小时check_metrics + 每天check_goals(AI评估)',
        '告警升级: info 24h未读→warning, warning 72h→critical',
        'WebSocket推送: ws_manager.broadcast_to_workspace()'
    ]),
    ('L10 战略决策', [
        'Decision Agent: CSO角色, 5维度分析框架(市场/财务/运营/客户/风险)',
        '上下文: Company Profile+Memory+Metrics三层聚合→LLM',
        '持久化: decisions表存储, 历史可查询/展开/删除',
        '独立于Chat: 专用端点, 不污染对话历史'
    ]),
]
for name, details in layers:
    h2(name)
    for d in details: bullet(d)

# ═══ Ch4: Agent System ═══
h1('4. Multi-Agent 系统')
h2('4.1 架构')
para('LangGraph StateGraph + Supervisor Pattern。16个Agent统一编排与路由。')
h2('4.2 执行流程')
for s in ['1. 用户输入→WebSocket/SSE发送到后端',
          '2. AgentOrchestrator.run_stream()启动流水线',
          '3. Supervisor Node分析→LLM路由(pipeline)或Keyword Fallback',
          '4. 上下文预装: RAG→Web→Memory→Profile→Metrics→Documents',
          '5. Agent流水线顺序执行: node_fn(state)→每步输出传给下一个',
          '6. 流式返回: token→WebSocket→前端逐字渲染',
          '7. 对话后: 自动记忆提取(LLM→JSON→pgvector) + 报告保存']:
    bullet(s)

h2('4.3 Agent实现模式')
for p in ['专属System Prompt(角色+分析框架+输出格式)',
          'run_X_agent(query,context_text)→{"final_response":str,"agent_trace":list}',
          'Supervisor Node函数→StateGraph.add_node()',
          'Orchestrator._ensure_nodes()注册']:
    bullet(p)

h2('4.4 上下文注入')
for s in ['Step 5: RAG检索(RagService.search→5条)',
          'Step 5.5: 记忆召回(MemoryService)',
          'Step 5.6: 企业画像(CompanyService)',
          'Step 5.7: 文档列表+文件名匹配',
          'Step 5.8: 数字孪生指标(BusinessMetricsService)']:
    bullet(s)

# ═══ Ch5: Data ═══
h1('5. 数据架构')
h2('5.1 存储分布')
for k,v in [('PostgreSQL+pgvector','22张业务表+向量嵌入(512维)'),
            ('Neo4j','知识图谱,7种实体+8种关系'),
            ('Redis','缓存+Celery消息队列+WebSocket Pub/Sub'),
            ('文件系统','上传文档/app/uploads/,模型缓存/app/model_cache/')]:
    bullet(f'{k}: {v}')

h2('5.2 数据关系')
para('Companies(1:1 Workspace)→ Departments(树形自引用)→ Positions→ Employees')
para('BusinessMetrics(时序数据)→ Predictions(预测结果)')
para('EnterpriseMemories(向量嵌入)→ MemoryEvents')
para('ProactiveAlerts(告警) | Decisions(战略决策)')

# ═══ Ch6: API ═══
h1('6. API 架构')
h2('6.1 路由结构')
routes = [
    '认证: /api/v1/auth (login/register/refresh)',
    '企业画像: /api/v1/workspaces/{id}/company (37端点,9实体CRUD)',
    '文档: /api/v1/workspaces/{id}/documents (upload/list/preview/delete)',
    '知识: /api/v1/workspaces/{id}/knowledge (search/connect/transcribe)',
    '记忆: /api/v1/workspaces/{id}/memories (8端点)',
    '图谱: /api/v1/workspaces/{id}/graph (query/search/stats/sync)',
    '指标: /api/v1/workspaces/{id}/metrics (record/batch/snapshot/trend)',
    '分析: /api/v1/workspaces/{id}/analytics (dashboard/analyze/alerts)',
    '预测: /api/v1/workspaces/{id}/forecast (predict/forecast/dashboard)',
    '告警: /api/v1/workspaces/{id}/alerts/proactive (check/dismiss)',
    '决策: /api/v1/workspaces/{id}/chat/decision (analyze/history/delete)',
    '对话: /api/v1/workspaces/{id}/chat (sessions/send/stream)',
]
for r in routes: bullet(r)

h2('6.2 认证与权限')
para('JWT Token: 登录→access_token(60min)+refresh_token(7d)。Axios拦截器自动401刷新。RBAC: Admin(全部)/Member(CRUD)/Viewer(只读)。')

# ═══ Ch7: Frontend ═══
h1('7. 前端架构')
h2('7.1 页面路由(16页)')
for p in ['/chat - 多Agent对话(WebSocket流式)',
          '/documents - 文档上传/预览/删除',
          '/knowledge - 语义搜索+6连接器+音频',
          '/graph - 知识图谱统计/搜索/同步',
          '/memories - 记忆+事件时间线+召回',
          '/analytics - KPI仪表盘+趋势图+AI分析+Alerts',
          '/forecast - 预测引擎(移动平均/线性回归)',
          '/decision-center - 战略决策中心(SSE流式)',
          '/reports - 报告列表与下载',
          '/members - 成员管理(搜索/添加/角色)',
          '/settings/company - 企业基本信息',
          '/settings/org - 组织架构(部门/职位/员工)',
          '/settings/business - 业务管理(产品/客户/流程)',
          '/settings/goals - Goals & KPIs',
          '/settings/metrics - 业务指标录入(批量JSON)',
          '/login - 登录/注册']:
    bullet(p)

h2('7.2 状态管理')
para('Zustand 3个Store: useAuthStore(token+user+login/logout)、useWorkspaceStore(activeWorkspaceId)、useChatStore(messages+streaming+session)。Token持久化到localStorage。')

h2('7.3 WebSocket')
para('自定义WebSocketClient: 自动重连(指数退避3s-30s,10次)、30s ping心跳、事件驱动(on/off)。Chat优先WebSocket，失败时回退到HTTP SSE。')

# ═══ Ch8: Deploy ═══
h1('8. 部署架构')
para('6个Docker服务通过docker-compose.yml编排:')
for s in ['PostgreSQL(pgvector/pg16) - 关系数据+向量存储',
          'Redis(7-alpine) - 缓存+Celery消息代理',
          'Neo4j(5-community) - 知识图谱',
          'Backend(Python:3.12-slim) - FastAPI+Celery+Alembic自动迁移',
          'Frontend(Node:22-alpine) - Next.js standalone模式',
          'Nginx(alpine) - 反向代理,统一入口:80']:
    bullet(s)

# ═══ Ch9: Flow ═══
h1('9. 系统运行流程')
h2('9.1 完整用户旅程')
journey = [
    '1. 注册/登录→JWT Token→前端Zustand存储',
    '2. 创建Workspace→数据隔离容器',
    '3. Settings页面录入Company/Departments/Employees/Products/Goals/KPIs',
    '4. Metrics页面手动录入或批量JSON导入→business_metrics表',
    '5. 文档上传→后端解析→Chunk→Embedding→pgvector',
    '6. Chat页面输入→Supervisor路由→Agent流水线→流式返回',
    '7. Analytics页面自动加载→KPI卡片+趋势图+AI分析+Alerts',
    '8. Forecast页面预测→移动平均/线性回归→图表渲染',
    '9. Decision Center输入问题→SSE流式5维度分析→结果持久化',
]
for step in journey: bullet(step)

h2('9.2 数据流')
para('用户操作→前端API调用(axios→/api/v1/...)→FastAPI端点(Depends注入: auth→role check→db session)→Service层→Model层(ORM→PostgreSQL/Neo4j)→响应(JSON/SSE/WebSocket)→前端渲染')

h2('9.3 Agent对话流')
para('用户消息(WebSocket)→AgentOrchestrator.run_stream()→Supervisor路由→上下文预装(6层)→Agent流水线顺序执行→流式token返回→保存回复→自动提取记忆')

# ═══ Ch10: Appendix ═══
h1('10. 附录')
h2('10.1 完整Agent列表')
for a in ['chat(通用)','search(RAG检索)','research(深度研究)','analyst(数据分析)',
          'writer(报告)','sql(数据库查询)','profile(企业画像)','memory(记忆)',
          'ceo(战略)','finance(财务)','sales(销售)','hr(人力)',
          'customer(客户)','operations(运营)','procurement(采购)','decision(决策)']:
    bullet(a)

h2('10.2 关键技术决策')
for d in ['DeepSeek API主力LLM,OpenAI兼容接口可替换',
          'pgvector替代独立向量数据库,减少运维复杂度',
          'Neo4j Community Edition,满足中小企业图谱需求',
          'Celery+Redis而非Kafka,降低部署门槛',
          'Zustand替代Redux,轻量级状态管理',
          'Recharts替代Plotly/Superset,零额外依赖',
          '自定义WebSocketClient替代socket.io,精简包体积']:
    bullet(d)

doc.save('EAIOS_Tech_Doc.docx')
print('Saved: EAIOS_Tech_Doc.docx')
