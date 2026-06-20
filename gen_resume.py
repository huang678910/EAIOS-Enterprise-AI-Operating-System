"""生成简历项目经历（多版本）"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.size = Pt(10.5)
style.font.name = 'Microsoft YaHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def h1(text):
    doc.add_page_break()
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.size = Pt(15)
    return h

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.size = Pt(12)
    return h

def body(text):
    doc.add_paragraph(text)

def resume_block(lines):
    """简历风格的条目"""
    for line in lines:
        p = doc.add_paragraph(line, style='List Bullet')
        for r in p.runs: r.font.size = Pt(10.5)

def code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def tip(text):
    p = doc.add_paragraph()
    r = p.add_run(f"💡 {text}")
    r.font.size = Pt(9); r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return p

# ═══ 封面 ═══
for _ in range(2): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('简历项目经历 · EAIOS')
r.font.size = Pt(22); r.bold = True; r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

doc.add_paragraph()
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f'{datetime.date.today().strftime("%Y-%m-%d")}  |  多版本·可直接复制到简历').font.size = Pt(10)

# ═══ 版本 1：一句话 ═══
h1('版本 1：一句话概述（自我介绍/简历头部）')

body('以下选一句最适合你投递岗位的：')

h2('侧重 Agent 编排（推荐）')
code_block('独立设计并实现了一个基于 LangGraph + Supervisor Pattern 的 16-Agent Multi-Agent 企业智能操作系统，涵盖 RAG 检索增强生成、知识图谱、数字孪生、预测引擎与主动告警，完成从数据录入到智能决策的全链路闭环。')

h2('侧重全栈工程')
code_block('全栈独立开发 EAIOS 企业 AI 操作系统：FastAPI + Next.js 15 + LangGraph Multi-Agent 编排，22 张数据表、19 个 API 模块、16 个前端页面、Docker Compose 6 服务编排部署。')

h2('侧重 LLM/RAG')
code_block('独立完成基于 pgvector + BGE Embedding 的 RAG 知识库系统与 LangGraph 驱动的 Multi-Agent 协作框架，支持 10 格式文档解析、向量语义检索、WebSocket 流式对话，Agent 自动感知企业上下文。')

# ═══ 版本 2：简历标准格式 ═══
h1('版本 2：简历标准条目（3-5 条 Bullet）')

h2('投递 AI Agent / LLM 应用岗位')
resume_block([
    '基于 LangGraph StateGraph + Supervisor Pattern 构建 16-Agent Multi-Agent 编排系统，实现 LLM 路由 + Keyword Fallback 双重路由策略，支持单 Agent 执行和多 Agent 流水线（如 search→analyst→writer 三段式）',
    '设计 6 层上下文自动注入机制（RAG 检索 / 企业记忆 / 企业画像 / 文档列表 / 数字孪生），Agent 在执行前自动获得企业全貌感知',
    '实现完整 RAG 流水线：10+ 格式文档解析（PyMuPDF/python-docx/Whisper 音频）→ tiktoken 分块 → BGE Embedding（512 维）→ pgvector 向量检索 → Query Rewrite → LLM 合成答案',
    '集成 pgvector 向量检索 + Neo4j 知识图谱双存储：向量搜索处理非结构化文档问答，Cypher 图查询处理"谁在哪个部门"类关系推理，分词搜索实现模糊实体匹配',
    '自定义 WebSocket 流式客户端（指数退避重连 + Ping 心跳 + SSE 降级），支持 token 级流式渲染；实现 JWT 双 Token 机制 + axios 拦截器自动刷新 + 并发防抖',
])

h2('投递全栈 / 后端岗位')
resume_block([
    '独立设计并实现企业 AI 操作系统全栈项目：后端 FastAPI + SQLAlchemy 2.0 Async + Celery + Redis，前端 Next.js 15 + React 19 + TypeScript + TailwindCSS',
    '设计 22 张数据表（PostgreSQL + pgvector）+ 13 个 Alembic 迁移，涵盖企业画像、文档知识库、企业记忆、数字孪生、预测、告警、决策 7 大业务域',
    '搭建 Neo4j 知识图谱（7 实体 × 8 关系），Cypher 白名单安全查询 + 分词搜索；Redis 三重角色：Celery Broker + WebSocket Pub/Sub + 通用缓存',
    '实现定时告警系统：Celery Beat 调度（每小时指标检测 + 每天目标检测）+ 三层检测架构（规则引擎→统计异常→AI 评估）+ WebSocket 实时推送 + 告警升级与去重',
    'Docker Compose 6 服务编排（Postgres/Redis/Neo4j/Backend/Frontend/Nginx），多阶段 Docker 构建优化，镜像体积压缩 60%+',
])

h2('投递前端岗位')
resume_block([
    '使用 Next.js 15 App Router + React 19 + TypeScript 构建 16 个前端页面，包含 Chat（WebSocket 流式对话）、Analytics（KPI 仪表盘）、Forecast（预测引擎）、Decision Center（SSE 流式分析）等',
    '基于 shadcn/ui + TailwindCSS 构建组件体系，实现响应式布局（移动端 Hamburger 菜单 + 侧边栏 Overlay）、Skeleton 骨架屏、Error Boundary 错误边界',
    'Zustand 三 Store 状态管理（Auth/Workspace/Chat），自定义 WebSocketClient 类（事件驱动 + 自动重连 + Ping 心跳 + SSE 回退）',
    '实现 Recharts 数据可视化（趋势图/进度条/告警面板）、SSE 流式 Markdown 渲染、批量 JSON 导入等交互功能',
])

# ═══ 版本 3：项目描述 ═══
h1('版本 3：完整项目描述（适合放 Portfolio/作品集）')

body('EAIOS（Enterprise AI Operating System）是一个面向企业的 AI 操作系统，核心通过 Multi-Agent 协作完成从数据管理到智能决策的闭环。')
body('')
body('架构亮点：')
resume_block([
    '16 个角色化 Agent 在 LangGraph StateGraph 下通过 Supervisor Pattern 统一编排，Agent 间共享 TypedDict State 传递上下文，支持 LLM 智能路由 + Keyword Fallback 降级',
    '6 层上下文自动注入：RAG 检索 → 记忆召回 → 企业画像 → 文档列表 → 数字孪生指标，Agent 执行前自动获取企业全貌',
    '10 层功能架构：企业画像(L1)→知识中心(L2)→企业记忆(L3)→知识图谱(L4)→数字孪生(L5)→部门 Agent(L6)→分析中心(L7)→预测引擎(L8)→主动式 AI(L9)→战略决策(L10)',
    'RAG 流水线：10 格式文档解析 → tiktoken 512 分块 → BGE Embedding 向量化 → pgvector 余弦检索 → Query Rewrite → LLM 合成带引用答案',
    'Proactive Alerts 三层检测：规则引擎(KPI 阈值) → 统计异常(Z-score/环比/反转) → AI 评估(LLM 过滤)，Celery Beat 定时调度 + WebSocket 实时推送 + 告警升级',
])
body('')
body('技术栈：FastAPI + SQLAlchemy 2.0 Async + LangGraph + LangChain + DeepSeek API + PostgreSQL/pgvector + Neo4j + Redis + Celery + Next.js 15 + TypeScript + Docker Compose')
body('项目规模：~20,000 行代码 / 22 张数据表 / 13 个数据库迁移 / 19 个 API 路由模块 / 100+ REST 端点 / 6 Docker 服务')

# ═══ 版本 4：STAR 故事 ═══
h1('版本 4：STAR 法则面试故事（行为面试用）')

h2('故事 1：从零搭建 Multi-Agent 系统')
body('S（情境）：项目需要让用户能用自然语言与企业数据交互，但单 LLM 无法感知企业个性化数据，也无法处理复杂的多步骤任务。')
body('T（任务）：设计一个 Multi-Agent 系统，让不同 Agent 各司其职、协作完成任务，同时每个 Agent 能自动获取相关企业上下文。')
body('A（行动）：选用 LangGraph StateGraph 做编排框架，设计 Supervisor Pattern（一个 Supervisor Agent 负责路由，16 个 Worker Agent 各负责一个领域）。实现了 6 层上下文注入机制，在 Agent 执行前自动从 RAG、记忆、画像、指标等数据源拉取上下文。LLM 路由失败时自动降级为关键词匹配。')
body('R（结果）：系统支持 16 个 Agent 的灵活编排，用户可以用自然语言完成文档检索、数据分析、财务诊断、战略决策等多样化任务。WebSocket 流式返回让对话体验接近 ChatGPT。')

h2('故事 2：解决 Docker 部署的依赖地狱')
body('S（情境）：项目需要 Docker 化部署，但 pip install 下载超时（30 分钟+），镜像体积超过 5GB。')
body('T（任务）：优化 Docker 构建流程，确保镜像能在合理时间内构建完成，且体积可控。')
body('A（行动）：分析依赖链发现 sentence-transformers → torch → 10+ NVIDIA CUDA 包（1.5GB+）。在 Dockerfile 中先安装 CPU-only PyTorch 再安装其他依赖，阻止 CUDA 包下载。同时配置国内镜像源（阿里云 Debian + PyPI），修复 passlib + bcrypt 版本兼容性问题。')
body('R（结果）：构建时间从 30 分钟+降至 10 分钟内，镜像体积压缩 60%。项目可通过 docker-compose up -d 一键部署。')

# ═══ 保存 ═══
doc.save('EAIOS_Resume_Experience.docx')
print('Saved: EAIOS_Resume_Experience.docx')