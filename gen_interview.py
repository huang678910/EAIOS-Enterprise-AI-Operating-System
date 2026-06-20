"""生成 EAIOS 全方面面试题 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.size = Pt(10.5)
style.font.name = 'Microsoft YaHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def h1(text):
    doc.add_page_break()
    return doc.add_heading(text, level=1)

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    return p

def question(q, difficulty="★★☆"):
    p = doc.add_paragraph()
    r = p.add_run(f"\n【{difficulty}】{q}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    return p

def answer(a):
    p = doc.add_paragraph()
    r = p.add_run(a)
    r.font.size = Pt(10)
    return p

def code_block(code):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

# ══════════════════════════════════════════
# 封面
# ══════════════════════════════════════════
for _ in range(3): doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('EAIOS — 企业智能操作系统')
r.font.size = Pt(28); r.bold = True; r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('全方面技术面试题（含参考答案）')
r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'v1.0  |  {datetime.date.today().strftime("%Y-%m-%d")}  |  共 7 大类 / 30 题').font.size = Pt(11)

doc.add_paragraph()
desc = doc.add_paragraph(); desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc.add_run('基于项目实际代码（16 Agent / 19 API / 16 页面 / 22 数据表 / 6 Docker 服务）').font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════
# 目录
# ══════════════════════════════════════════
h1('目录')
toc_items = [
    '一、系统架构设计（5 题）',
    '二、Multi-Agent 系统（5 题）',
    '三、后端开发 — FastAPI + 数据库（5 题）',
    '四、前端开发 — Next.js + React（4 题）',
    '五、RAG & 知识库 & 图谱（4 题）',
    '六、数字孪生 & 预测 & 告警（4 题）',
    '七、部署 & DevOps（3 题）',
]
for item in toc_items:
    para(item)

# ══════════════════════════════════════════
# 第一章：系统架构设计
# ══════════════════════════════════════════
h1('一、系统架构设计（5 题）')

question('1. 请描述 EAIOS 的 10 层架构，并说明每层的核心职责。', '★★★★★')
answer("""EAIOS 采用 10 层渐进式企业智能架构：

L1 企业画像（Company Profile）：9 张核心表（companies/departments/positions/employees/products/customers/processes/goals/kpis），37 个 REST 端点。录入企业基本信息、组织架构、产品/客户/流程、Goals & KPIs，是整个系统的"地基数据"。

L2 知识中心（Knowledge Hub）：文档解析流水线（10+ 格式）→ 分块（512 token）→ BGE Embedding（512 维）→ pgvector 存储。6 个外部连接器（GitHub/Notion/Jira/Confluence/Slack/WeCom），支持 Whisper 音频转录。QA 模式：Query Rewrite（LLM）→ pgvector 余弦搜索 → Reranker → LLM 合成答案。

L3 企业记忆（Enterprise Memory）：三种记忆类型：长期记忆（Long-Term）、事件记忆（Episodic）、语义记忆（Semantic）。对话后 LLM 自动判断是否值得记忆 → JSON → pgvector → 去重入库。混合加权召回：重要性×40% + 时效性×30% + 向量相似度×30%。

L4 知识图谱（Knowledge Graph）：Neo4j 存储 7 种实体（Company/Department/Employee/Product/Customer/Process/Goal）+ 8 种关系。Cypher 白名单安全机制（只允许 MATCH/RETURN/CALL）。分词搜索：tokenize 查询词 → 匹配任一实体属性。

L5 数字孪生（Digital Twin）：business_metrics 表存储时序业务数据。Agent 对话时自动注入指标上下文。Analytics 仪表盘聚合快照 + 趋势。预测引擎基于历史数据外推。

L6 部门 Agent（Department Agents）：7 个角色化 Agent（CEO/CFO/VP Sales/CHRO/COO/CPO/VP Customer），每个有专属 System Prompt + run_X_agent() 函数。Supervisor 负责路由。

L7 分析中心（Analytics Center）：聚合 metrics snapshot + trends + KPIs + Goals + alerts。业务优先级排序（revenue → cost → hr → ...）。AI 分析（LLM）缓存 1 小时。

L8 预测引擎（Forecast Engine）：加权移动平均 + 线性回归（sklearn）。自动推算预测期数 = min(历史数/2, 6)。置信区间 ±1.5×标准差。

L9 主动式 AI（Proactive AI）：三层检测：规则引擎（KPI 阈值）→ 统计异常（Z-score/环比/趋势反转）→ AI 评估（LLM 过滤）。Celery Beat 定时调度（每小时/每天）。告警升级：info 24h → warning，warning 72h → critical。WebSocket 实时推送。

L10 战略决策中心（Decision Center）：CSO 角色 Agent，5 维度分析框架（市场/财务/运营/客户/风险）。SSE 流式返回，结果持久化到 decisions 表，支持历史回溯。独立于 Chat 对话系统。""")

# ──
question('2. 系统使用了哪些数据存储？各自的职责是什么？为什么这样选型？', '★★★★★')
answer("""系统使用了 5 种数据存储：

1. PostgreSQL + pgvector：关系数据（22 张业务表）+ 向量嵌入（512 维 BGE）。选型原因：pgvector 扩展让关系库和向量库合二为一，减少运维复杂度，避免单独的 Milvus/Pinecone 集群。

2. Neo4j 5 Community：知识图谱存储。7 种实体 + 8 种关系，支持 Cypher 查询。选 Community Edition 是因为中小企业场景不需要企业版的高级集群功能，单机足以处理万级节点。

3. Redis：三重角色：
   - Celery 消息代理（broker + backend）
   - WebSocket Pub/Sub（频道广播）
   - 通用缓存（AI 分析结果 TTL 1 小时）

4. 文件系统：上传文档存储（/app/uploads/），HuggingFace 模型缓存（/app/model_cache/）。

5. Docker Volumes：pgdata/redisdata/neo4j_data/uploads/model_cache 六个命名卷，保证容器重启数据不丢失。

选型原则：在满足需求的前提下最小化组件数量，PostgreSQL + pgvector 替代了"PostgreSQL + 独立向量数据库"的传统方案。""")

# ──
question('3. 请画出系统在 Docker Compose 下的部署拓扑，并说明 Nginx 的职责。', '★★★★☆')
answer("""6 服务拓扑：

┌─────────────────────────────────────────────┐
│                  Nginx (:8080)               │
│  / → frontend:3000   /api/* → backend:8000  │
│         /ws/* → backend:8000 (Upgrade)       │
└──────┬────────────────────┬─────────────────┘
       │                    │
┌──────▼──────┐    ┌───────▼───────┐
│  Frontend   │    │   Backend     │
│  Node:22    │    │ Python:3.12  │
│  Next.js 15 │    │  FastAPI      │
│  :3000      │    │  :8000        │
└─────────────┘    └───┬───┬───┬───┘
                       │   │   │
          ┌────────────┼───┼───┼────────┐
          │            │   │   │        │
     ┌────▼───┐  ┌────▼───▼┐ ┌▼──────┐ ┌▼──────┐
     │Postgres│  │  Redis  │ │ Neo4j │ │  FS   │
     │pgvector│  │ :6379   │ │:7474  │ │uploads│
     │:5432   │  │(broker) │ │:7687  │ │       │
     └────────┘  └─────────┘ └───────┘ └───────┘

Nginx 职责：
1. 反向代理：统一入口，将 / 路由到前端，/api/* 路由到后端
2. WebSocket 代理：/ws/* 升级为 WebSocket 连接
3. CORS 处理：在代理层统一处理跨域
4. 静态资源服务：gzip 压缩，缓存头控制
5. 安全性：隐藏后端端口，限制请求大小""")

# ──
question('4. 前端 API 请求在开发环境和生产环境分别是如何路由的？为什么要这样设计？', '★★★☆☆')
answer("""开发环境：
- Next.js dev server 运行在 localhost:3000
- FastAPI 运行在 localhost:8000
- 前端 .env.local 设置 NEXT_PUBLIC_API_URL=http://localhost:8000
- 浏览器直接跨域请求后端

生产环境（Docker）：
- 前端 .env.local 设置 NEXT_PUBLIC_API_URL=（空字符串）
- API_BASE = ""，axios 发起相对路径请求（如 /api/v1/auth/login）
- 浏览器请求发送到同源（localhost:8080）
- Nginx 将 /api/* 代理到 backend:8000

设计原因：
- 开发时前端 3000 端口访问，后端 8000 端口，需要显式指定跨域 URL
- 生产时所有请求经 Nginx 统一入口，相对路径消除了环境差异
- 相对路径方案让同一份构建产物在任何域名下都能正常工作
- 避免在 JS bundle 中硬编码绝对 URL，符合 12-Factor App 原则""")

# ──
question('5. 系统的认证和权限模型是怎样的？JWT Token 过期后如何处理？', '★★★☆☆')
answer("""认证流程：
1. POST /api/v1/auth/login → 后端验证密码 → 返回 access_token（60 分钟）+ refresh_token（7 天）
2. 前端 axios 拦截器自动在请求头附加 Authorization: Bearer <token>
3. 后端 Depends(get_current_user) 解析 JWT，注入 User 对象

权限模型（RBAC 3 级）：
- Admin：全部操作（创建/修改/删除 + 成员管理）
- Member：CRUD 操作
- Viewer：只读
- require_workspace_role(workspace_id, user, "member", db) 在每个端点入口检查

Token 自动刷新机制：
1. axios 响应拦截器检测 401
2. 用 refresh_token 调用 POST /api/v1/auth/refresh
3. 成功 → 更新 localStorage 中的 token，重放原始请求
4. 刷新时加锁（isRefreshing=true），并发请求排队等候，避免多次刷新
5. 失败 → 清除 token，重定向到 /login""")

# ══════════════════════════════════════════
# 第二章：Multi-Agent 系统
# ══════════════════════════════════════════
h1('二、Multi-Agent 系统（5 题）')

question('1. 请详细描述 Supervisor Pattern 在项目中的实现：路由决策、流水线执行、上下文注入的全流程。', '★★★★★')
answer("""项目使用 LangGraph StateGraph + Supervisor Pattern 实现 16 Agent 编排。

路由决策（Supervisor Node）：
1. 用户消息 → AgentOrchestrator.run_stream()
2. Supervisor Node 调用 LLM（DeepSeek），输入 ROUTE_SYSTEM_PROMPT（含所有 Agent 能力描述 + 流水线规则）
3. LLM 返回 JSON：{"next_agents": ["search","finance"], "reason": "..."}
4. Keyword Fallback：如果 LLM 路由失败，用关键词匹配兜底

Agent 路由规则示例：
- "我们公司上季度收入多少" → ["profile","finance"]（先查画像，再分析财务）
- "搜索文档中的销售数据并写报告" → ["search","analyst","writer"]
- "简单问候" → ["chat"]

流水线执行（Pipeline）：
1. supervisor_node() 的输出 state["next_agents"] 是一个有序列表
2. LangGraph 按顺序执行：agent_1(state) → agent_2(state) → ... → agent_n(state)
3. 每个 Agent 的输出保存在 state["agent_trace"] 中
4. 前一个 Agent 的上下文自动传给下一个 Agent
5. 流式返回：每个 token 通过 WebSocket yield 到前端

上下文注入（6 层预处理）：
Step 5: RAG 检索（RagService.search → 5 条相关文档）
Step 5.5: 记忆召回（MemoryService）
Step 5.6: 企业画像（CompanyService）
Step 5.7: 文档列表 + 文件名匹配
Step 5.8: 数字孪生指标（BusinessMetricsService）
Step 6: 最终 context_text 注入 Agent System Prompt

AgentState（TypedDict）结构：
- messages: 对话历史
- user_query: 用户问题
- next_agents: 路由决定的 Agent 列表
- context_text: 预装的上下文
- sources: RAG 检索到的文档
- final_response: Agent 输出
- agent_trace: Agent 执行链""")

# ──
question('2. 7 个部门 Agent 的设计模式是什么？如果让你新增一个"法务 Agent"，需要改哪些文件？', '★★★★☆')
answer("""部门 Agent 设计模式：

1. 专属 System Prompt：定义角色 + 分析框架 + 输出格式
   示例（CFO Agent）：
   "You are the CFO of this company. Analyze financial data including
    revenue, costs, margins, cashflow. Consider: 1. Revenue trends
    2. Cost structure 3. Profitability 4. Financial risks 5. Recommendations"

2. run_X_agent(query, context_text) 函数签名：
   async def run_finance_agent(query: str, context_text: str) -> dict:
       prompt = ChatPromptTemplate.from_messages([
           ("system", CFO_SYSTEM_PROMPT),
           ("system", "Company Context:\n{context}"),
           ("human", "{query}")
       ])
       # 调用 LLM，流式输出
       return {"final_response": response, "agent_trace": [...]}

3. LangGraph Node 函数：
   async def finance_node(state: AgentState) -> AgentState:
       result = await run_finance_agent(
           state["user_query"], state.get("context_text", "")
       )
       state["final_response"] = result["final_response"]
       state["agent_trace"].append("finance")
       return state

4. 注册到 Supervisor：在 supervisor.py 中更新 AGENT_NAMES 和 ROUTE_SYSTEM_PROMPT

新增"法务 Agent"需要修改的文件：
1. 新建 backend/app/agents/legal_agent.py — 定义 SYSTEM_PROMPT + run_legal_agent()
2. 修改 backend/app/agents/supervisor.py — 在 ROUTE_SYSTEM_PROMPT 中加入 legal Agent 能力描述
3. 修改 backend/app/services/agent_orchestrator.py — _ensure_nodes() 中导入 legal_node 并注册
4. 可选：在 Keyword Fallback 路由中加入法务相关关键词""")

# ──
question('3. AgentOrchestrator.run_stream() 是如何实现流式返回的？WebSocket 和 SSE 分别用在什么场景？', '★★★★☆')
answer("""流式返回机制：

AgentOrchestrator.run_stream() 是一个 AsyncGenerator：
   async def run_stream(self, session_id, message) -> AsyncGenerator[dict, None]:
       # 1. 保存用户消息
       # 2. 构建 AgentState
       # 3. Supervisor 路由
       yield {"type": "status", "content": "Routing...", "agent": "supervisor"}

       # 4. 上下文预装（6 步）
       yield {"type": "status", "content": "Loading context...", "agent": "search"}

       # 5. 执行 Agent 流水线
       for agent_name in state["next_agents"]:
           node_fn = AGENT_NODES[agent_name]
           async for chunk in node_fn(state):
               yield chunk  # token-level streaming

       # 6. 保存回复
       yield {"type": "done", "content": state["final_response"]}

WebSocket 使用场景（Chat 页面）：
- 建立 ws://localhost:8080/api/v1/ws/chat?token=xxx&session_id=xxx
- 双向实时通信，支持 token 级流式输出
- 自定义 WebSocketClient，带指数退避自动重连（3s-30s，最多 10 次）
- 30 秒 ping 心跳保持连接

SSE 使用场景（Decision Center）：
- POST /api/v1/workspaces/{id}/chat/decision/analyze
- 单向流式（服务器 → 客户端），适合决策分析的长文本生成
- 基于 sse-starlette 库实现
- 通过 fetch + ReadableStream 在前端消费""")

# ──
question('4. LLM 路由失败时的 Keyword Fallback 机制是如何实现的？为什么需要它？', '★★★☆☆')
answer("""实现原理：

每个 Agent 在 supervisor.py 中定义了一组触发关键词：
   KEYWORD_ROUTES = {
       "finance": ["收入","成本","利润","财务","预算","cashflow","ROI"],
       "sales": ["销售","客户获取","增长","定价","pipeline"],
       "hr": ["招聘","人力","组织","headcount"],
       "profile": ["公司","部门","产品","员工","企业"],
       ...
   }

Fallback 逻辑：
1. LLM 路由正常时，直接使用 LLM 返回的 next_agents
2. 如果 LLM 调用抛异常（超时/API 错误/JSON 解析失败），触发 fallback
3. Fallback 算法：遍历用户 query，匹配每个 Agent 的关键词
4. 命中 ≥1 个关键词 → 加入路由列表
5. 如果一个都没命中 → 默认路由到 ["chat"]

为什么需要？
- LLM API 可能不可用（网络波动、额度耗尽）
- 简单问题不需要 LLM 路由（如"列出部门"→ 直接 profile Agent）
- 降低延迟：关键词匹配是 O(n) 字符串查找，远快于 LLM 推理
- 提高可靠性：双重路由策略确保系统在 LLM 故障时仍然可用""")

# ──
question('5. 对话结束后系统自动做了什么？记忆提取的具体流程是什么？', '★★★☆☆')
answer("""对话后自动处理流程：

1. 保存消息：ChatService.save_message() 将 user + assistant 消息写入 messages 表
2. 记忆提取（自动触发）：
   a. LLM 判断：给 LLM 完整对话上下文，问"这段对话中有值得记住的事实/决定/偏好吗？"
   b. 如果 LLM 返回有意义的 JSON：{"facts": ["公司Q2收入目标为500万", "CEO偏好风险规避策略"]}
   c. 对每个 fact 生成 BGE Embedding（512 维）
   d. 去重检查：pgvector 余弦相似度 > 0.9 则跳过
   e. 写入 enterprise_memories 表：
      - type: long_term（事实）/ episodic（事件）/ semantic（语义）
      - importance: 1-10（LLM 评估）
      - last_accessed: 当前时间
      - embedding: pgvector 向量
   f. 同时写入 memory_events 表作为事件日志

3. 记忆衰减：
   - 30 天未被访问 → importance -0.5
   - 混合加权召回公式：importance × 40% + recency × 30% + vector_similarity × 30%""")

# ══════════════════════════════════════════
# 第三章：后端开发
# ══════════════════════════════════════════
h1('三、后端开发 — FastAPI + 数据库（5 题）')

question('1. FastAPI 的依赖注入在这个项目中是如何使用的？举 3 个例子。', '★★★★☆')
answer("""三个核心依赖注入模式：

1. 数据库 Session（Depends(get_db)）：
   async def get_db():
       async with AsyncSessionLocal() as session:
           yield session
   # 每个请求自动获取一个异步数据库会话，请求结束后自动关闭

2. 认证 + 用户注入（Depends(get_current_user)）：
   async def get_current_user(
       token: str = Depends(oauth2_scheme),
       db: AsyncSession = Depends(get_db)
   ) -> User:
       payload = jwt.decode(token, SECRET_KEY)
       user = await db.get(User, payload["sub"])
       return user
   # 端点自动获取当前登录用户，无需手动解析 Token

3. 权限检查（require_workspace_role）：
   async def require_workspace_role(
       workspace_id: str, user: User, min_role: str, db: AsyncSession
   ):
       member = await db.execute(
           select(WorkspaceMember).where(
               WorkspaceMember.workspace_id == workspace_id,
               WorkspaceMember.user_id == user.id
           )
       )
       if not member or ROLE_ORDER[member.role] < ROLE_ORDER[min_role]:
           raise HTTPException(403, "Insufficient permissions")

链式调用示例：
   @router.delete("/{metric_id}", status_code=204)
   async def delete_metric(
       workspace_id: UUID,
       metric_id: UUID,
       db: AsyncSession = Depends(get_db),
       current_user: User = Depends(get_current_user),
   ):
       await require_workspace_role(str(workspace_id), current_user, "admin", db)
       # ... 执行删除逻辑""")

# ──
question('2. SQLAlchemy 2.0 Async 在这个项目中如何使用？asyncpg 和 pgvector 分别解决什么问题？', '★★★★☆')
answer("""SQLAlchemy 2.0 Async 使用模式：

1. 声明式模型（Mapped + mapped_column）：
   class BusinessMetric(Base):
       __tablename__ = "business_metrics"
       id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), primary_key=True)
       metric_name: Mapped[str] = mapped_column(String(100), nullable=False)
       metric_value: Mapped[float] = mapped_column(Float, nullable=False)

2. 异步查询（AsyncSession）：
   result = await db.execute(
       select(BusinessMetric)
       .where(BusinessMetric.company_id == company.id)
       .order_by(BusinessMetric.category.asc().nulls_last())
   )
   metrics = result.scalars().all()

3. 异步引擎配置：
   engine = create_async_engine(
       f"postgresql+asyncpg://{user}:{pwd}@{host}:{port}/{db}",
       pool_size=20, max_overflow=10
   )

asyncpg 的作用：
- PostgreSQL 的高性能异步驱动，基于 asyncio 协议直接实现
- 比 psycopg2 快 2-3 倍（异步 I/O，无 GIL 限制）
- 支持连接池、预处理语句、COPY 协议批量导入

pgvector 的作用：
- PostgreSQL 扩展，新增 vector(512) 数据类型
- 支持余弦相似度（<=>）、欧氏距离（<->）、内积（<#>）运算符
- 支持 IVFFlat 索引加速近似搜索
- 让同一张表既存业务字段又存向量嵌入，无需独立向量数据库

实际查询示例：
   SELECT *, 1 - (embedding <=> query_embedding) AS similarity
   FROM enterprise_memories
   ORDER BY embedding <=> query_embedding
   LIMIT 10;""")

# ──
question('3. Alembic 数据库迁移在本项目中的工作流是怎样的？如何新增一张表？', '★★★☆☆')
answer("""Alembic 工作流：

1. 自动生成迁移脚本：
   alembic revision --autogenerate -m "add new table"

2. 查看生成的迁移文件（backend/alembic/versions/xxx_add_new_table.py）：
   def upgrade():
       op.create_table('new_table',
           sa.Column('id', UUID(as_uuid=True), primary_key=True),
           sa.Column('name', sa.String(200), nullable=False),
           ...
       )

   def downgrade():
       op.drop_table('new_table')

3. 应用迁移：
   alembic upgrade head

4. Docker 启动时自动迁移（Dockerfile CMD）：
   sh -c "alembic upgrade head && uvicorn app.main:app --host 0.0.0.0 --port 8000"

新增一张表的完整步骤：
1. 在 backend/app/models/ 下新建模型文件（或添加到已有文件）
2. 定义 SQLAlchemy Model 类
3. 在 backend/app/models/__init__.py 中 import 新模型
4. 运行 alembic revision --autogenerate -m "描述"
5. 检查生成的迁移脚本（确保 upgrade/downgrade 正确）
6. alembic upgrade head 或重建 Docker 镜像

项目目前有 13 个迁移文件，覆盖了从初始建表到后来逐步添加的
business_metrics、goal_direction、predictions、proactive_alerts、decisions 等表。""")

# ──
question('4. Celery + Celery Beat 在本项目中承担什么角色？定时任务是如何配置的？', '★★★☆☆')
answer("""Celery 在本项目中的角色：

1. 异步任务队列（基于 Redis broker）：
   - 文档解析任务（大文件解析不阻塞 API 响应）
   - 报告生成任务
   - 告警检查任务

2. Celery Beat 定时调度：

   # celery_app.py
   beat_schedule = {
       "check-metrics-hourly": {
           "task": "app.tasks.alert_tasks.check_metrics_alerts",
           "schedule": 3600.0,  # 每小时
       },
       "check-goals-daily": {
           "task": "app.tasks.alert_tasks.check_goals_alerts",
           "schedule": 86400.0,  # 每天
       },
   }

定时任务详解：
- check_metrics_alerts（每小时）：
  1. 获取所有 BusinessMetrics 最新快照
  2. 规则引擎：对比 KPI 阈值
  3. 统计异常检测：Z-score > 2.5 或环比 > 50% 或趋势反转
  4. 生成 ProactiveAlert 记录

- check_goals_alerts（每天）：
  1. 获取所有 CompanyGoal
  2. 计算 progress_pct
  3. 带 AI 评估（LLM 判断是否需要告警）
  4. 生成告警 + WebSocket 推送

启动方式：
- 开发环境：celery -A app.celery_app worker -l info & celery -A app.celery_app beat -l info
- Docker 环境：backend 容器中通过 uvicorn 启动（beat 集成在主进程）""")

# ──
question('5. 这个项目的错误处理模式是怎样的？请以"workspace 不存在"到"用户看到 404"为例说明完整流程。', '★★★☆☆')
answer("""错误处理完整链路：

1. 后端 Service 层（抛出业务异常）：
   async def _get_company(self, workspace_id: str) -> Company:
       result = await self.db.execute(...)
       company = result.scalar_one_or_none()
       if not company:
           raise ValueError("Company profile not found")

2. 后端 API 层（转换为 HTTP 异常）：
   try:
       svc = BusinessMetricsService(db)
       metrics = await svc.record(str(workspace_id), body)
   except ValueError as e:
       raise HTTPException(status_code=400, detail=str(e))

3. FastAPI 全局异常处理（未捕获的异常）：
   # 500 Internal Server Error → 统一 JSON 响应
   {"detail": "Internal Server Error"}

4. Nginx 层：
   # 后端不可达时返回 502 Bad Gateway
   # 超时时返回 504 Gateway Timeout

5. 前端 axios 拦截器：
   api.interceptors.response.use(
       (response) => response,  # 正常响应
       async (error) => {
           if (error.response?.status === 401) {
               // Token 刷新
           }
           // 其他错误 → 抛出给调用方
           return Promise.reject(error);
       }
   )

6. 前端页面：
   try {
       result = await login(email, password);
       router.push("/chat");
   } catch (err) {
       const detail = err.response?.data?.detail;
       setError(detail || "Authentication failed");
   }

7. 前端 Error Boundary（兜底）：
   - error.tsx：页面级错误边界，显示"出错了" + 重试按钮
   - global-error.tsx：根级别错误边界
   - not-found.tsx：自定义 404 页面（品牌风格）""")

# ══════════════════════════════════════════
# 第四章：前端开发
# ══════════════════════════════════════════
h1('四、前端开发 — Next.js + React（4 题）')

question('1. Next.js 15 App Router 在本项目中是如何组织路由的？请列举所有页面及对应路由。', '★★★★☆')
answer("""路由组织（Next.js 15 App Router，基于文件系统）：

Chat & 协作：
- /chat → src/app/chat/page.tsx（多 Agent 对话，WebSocket 流式）
- /documents → src/app/documents/page.tsx（文档上传/预览/删除）
- /reports → src/app/reports/page.tsx（报告列表）

知识 & 记忆：
- /knowledge → src/app/knowledge/page.tsx（语义搜索 + 6 连接器 + 音频转录）
- /graph → src/app/graph/page.tsx（知识图谱统计/搜索/同步）
- /memories → src/app/memories/page.tsx（记忆 + 事件时间线 + 召回）

分析 & 决策：
- /analytics → src/app/analytics/page.tsx（KPI 仪表盘 + 趋势图 + AI 分析 + Alerts）
- /forecast → src/app/forecast/page.tsx（预测引擎，移动平均/线性回归）
- /decision-center → src/app/decision-center/page.tsx（战略决策，SSE 流式）

设置：
- /settings/company → 企业基本信息
- /settings/org → 组织架构（部门/职位/员工）
- /settings/business → 业务管理（产品/客户/流程）
- /settings/goals → Goals & KPIs
- /settings/metrics → 业务指标录入

用户管理：
- /members → 成员管理（搜索/添加/角色）
- /login → 登录/注册

错误页面：
- not-found.tsx → 自定义 404
- error.tsx → 500 错误边界 + 重试
- global-error.tsx → 根级别错误边界""")

# ──
question('2. 项目的 Zustand 状态管理是如何设计的？useChatStore 的实现思路是什么？', '★★★☆☆')
answer("""Zustand 三个 Store：

1. useAuthStore：
   - token / refreshToken / user
   - setAuth(token, user) / logout()
   - Token 持久化到 localStorage
   - 页面刷新后从 localStorage 恢复

2. useWorkspaceStore：
   - activeWorkspaceId / workspaces
   - setActiveWorkspace(id) / setWorkspaces(list)

3. useChatStore：
   - messages: 消息列表
   - isStreaming: 是否正在接收流式响应
   - activeSessionId: 当前会话
   - sendMessage(text): 发送消息 + 建立 WebSocket
   - onReceive(chunk): 追加 token 到最新消息
   - 流式渲染：React 状态更新触发 UI 逐字显示

关键设计原则：
- 每个 Store 职责单一，互不耦合
- 使用 subscribe 而非复杂的状态派生
- WebSocket 客户端独立于 Store，通过事件回调更新状态
- 无 Redux 的样板代码（action/reducer/selector 分离）""")

# ──
question('3. 项目使用 shadcn/ui 组件库。请说明这个组件库的工作方式和优势，并列举项目中使用了哪些具体组件。', '★★★☆☆')
answer("""shadcn/ui 工作方式：

- 不是 npm 包，而是通过 npx shadcn-ui add 将源码复制到项目中
- 组件代码在 src/components/ui/ 下，完全可控可修改
- 基于 Radix UI（无障碍基元）+ TailwindCSS + class-variance-authority
- 组件导入：import { Button } from "@/components/ui/button"

优势：
1. 完全控制源码：可以修改任何组件实现，不依赖第三方包更新
2. 按需引入：只添加实际使用的组件，不会有不必要的包体积
3. Radix 基元：无障碍（ARIA）、键盘导航内置
4. TailwindCSS 集成：样式通过 className 覆盖，无需 CSS-in-JS

项目中使用的组件：
- Button, Input, Card（CardContent, CardHeader）
- Dialog, Sheet（侧边抽屉）
- Select, DropdownMenu
- Table, Badge, Tabs
- Skeleton（骨架屏，自定义扩展）
- toast / useToast""")

# ──
question('4. 自定义 WebSocketClient 是如何实现的？包含了哪些容错机制？', '★★★★☆')
answer("""WebSocketClient 实现（frontend/src/lib/websocket-client.ts）：

核心特性：

1. 自动重连（指数退避）：
   - 初始延迟 3 秒，每次乘以 2
   - 最大延迟 30 秒
   - 最多重试 10 次
   - 正常关闭（code 1000）不重连

2. Ping 心跳：
   - 每 30 秒发送 ping 消息
   - 收到 pong 重置超时计时器
   - 超时未收到 pong → 主动关闭 → 触发重连

3. 事件驱动架构：
   on(event, handler) / off(event, handler) / _emit(event)
   支持事件：connected, message, error, close, reconnecting

4. Token 集成：
   - 连接 URL 带 token 参数：ws://...?token=xxx
   - token 过期时服务端返回 4001 → 触发重新登录

5. 消息规范化（兼容新旧格式）：
   if (!msg.data && (msg.content || msg.sources)) {
       msg.data = { content: msg.content, sources: msg.sources, agent: msg.agent }
   }

6. Chat 优先 WebSocket，失败回退 HTTP SSE：
   try { ws.send(message) } catch { fallback to fetch SSE }
   这种双重策略确保即使 WebSocket 不可用，对话功能仍然能工作""")

# ══════════════════════════════════════════
# 第五章：RAG & 知识库 & 图谱
# ══════════════════════════════════════════
h1('五、RAG & 知识库 & 图谱（4 题）')

question('1. 请详细描述本项目的 RAG（检索增强生成）流水线：从文档上传到最终答案生成的完整过程。', '★★★★★')
answer("""RAG 流水线分为两个阶段：

═══ 离线阶段（文档入库）═══

1. 文档上传：POST /api/v1/workspaces/{id}/documents/upload
2. 格式解析（DocumentService）：
   - PDF → PyMuPDF (fitz) 提取文本 + 嵌入图片
   - DOCX → python-docx 提取段落
   - PPTX → python-pptx 提取文本 + 图片 + 表格
   - XLSX/CSV → pandas 读取，标记 content_type="table"
   - TXT/MD → 直接读取
   - 音频 → Whisper (small/base) 转写为文本
3. 分块（Chunking）：
   - 按 512 token 切割（tiktoken 计算）
   - 保留 50 token 重叠防止截断
4. 向量化：BGE-small-zh-v1.5（BAAI，512 维），匹配中文语义
5. 存储：pgvector INSERT，字段包含 document_id / chunk_index / content / embedding

═══ 在线阶段（问答）═══

1. Query Rewrite（查询重写）：
   LLM 将用户口语化问题改写为检索优化查询
   "咱们上季度赚了多少" → "上季度收入 利润 财务数据"

2. 向量检索：
   query_embedding = BGE.encode(rewritten_query)
   SELECT content, 1 - (embedding <=> query_embedding) AS score
   FROM document_chunks ORDER BY embedding <=> query_embedding LIMIT 5

3. Reranker 重排序（可选）：
   对候选 5 条结果按语义相关性二次排序

4. LLM 合成答案：
   System Prompt + RAG Context（top-5 chunks） + User Query
   → LLM 流式生成带引用的答案 → WebSocket 返回前端

5. 前端展示：
   - 引用来源标注（文档名 + chunk 位置）
   - 支持点击跳转到原始文档

RAG 上下文注入时机：
Orchestrator.run_stream() 的 Step 5，在 Agent 路由决策之后、Agent 执行之前""")

# ──
question('2. 你如何评估当前 BGE 嵌入模型的效果？如果检索结果不够好，有哪些优化方向？', '★★★★☆')
answer("""评估方法：
- 召回率@K：已知相关文档中，top-K 检索命中的比例
- MRR（Mean Reciprocal Rank）：第一个相关结果的排名的倒数，越接近 1 越好
- 人工评估：准备 20-30 个标准问题，人工判断 top-5 结果的相关性

优化方向（按性价比排序）：

1. Chunk 策略优化（零成本）：
   - 调整 chunk_size（当前 512 → 尝试 256/384/768）
   - 调整 overlap（当前 50 → 尝试 100/150）
   - 按文档结构分块（标题/段落边界），而非固定 token 数

2. Query Rewrite 优化（低成本）：
   - 在 Prompt 中增加领域术语表
   - 用历史问答对微调 Rewrite Prompt

3. 混合检索（Hybrid Search）：
   - 向量检索（语义）+ BM25（关键词）= 互补
   - 语义检索擅长同义词/近义表达，BM25 擅长精确匹配

4. 升级嵌入模型（中等成本）：
   - BGE-M3（1024 维，支持中英多语言）
   - BGE-large-zh-v1.5（1024 维，更高精度）
   - text2vec-large-chinese（1024 维）

5. Reranker 模型（高性价比）：
   - bge-reranker-v2-m3 对 top-20 候选重排序
   - 比增加 chunk 数量更有效

6. 元数据过滤（低成本）：
   - 按文档类型/上传时间/标签预过滤
   - 缩小检索范围提高精度""")

# ──
question('3. Knowledge Graph 和传统向量搜索的区别是什么？什么场景用 Graph 更合适？', '★★★★☆')
answer("""核心区别：

┌──────────────┬──────────────────┬─────────────────────┐
│    维度       │   向量搜索         │    知识图谱           │
├──────────────┼──────────────────┼─────────────────────┤
│ 原理          │ 向量余弦距离       │ 图遍历 (Cypher)      │
│ 优势          │ 语义理解/模糊匹配   │ 精确关系/多跳推理     │
│ 数据模型       │ 非结构化文本+向量   │ 实体-关系-实体（三元组）│
│ 查询方式       │ 自然语言→向量      │ Cypher 图查询        │
│ 典型问题       │ "怎么做..."       │ "X属于哪个部门"       │
└──────────────┴──────────────────┴─────────────────────┘

Graph 更适合的场景：
1. 关系查询："王经理的下属有哪些人？" → MATCH (e:Employee)-[:REPORTS_TO]->(m:Employee {name:'王经理'}) RETURN e
2. 多跳推理："销售部使用哪些产品服务哪些客户？" → 需要遍历 Department→Employee→Product→Customer 多跳
3. 层次结构："找出所有部门及其子部门" → 树形自引用依赖图遍历
4. 影响分析："如果换掉A产品的供应商，影响哪些部门？" → 多路径依赖遍历

向量搜索更适合的场景：
1. 文档问答："Q3营销策略是什么？" → 语义搜索文档库
2. 长文本理解：产品手册、合同条款等
3. 跨语言搜索：中文问题搜索英文文档

本项目中两者互补使用：
- Graph 回答结构化关系问题（谁在哪部门、产品归属）
- 向量搜索回答非结构化文档问答
- Agent 根据问题类型自动选择数据源""")

# ──
question('4. Neo4j 分词搜索的具体实现是什么？为什么要用分词而不是直接用 CONTAINS？', '★★★☆☆')
answer("""分词搜索实现（graph_service.py）：

def _search_entities(self, query: str):
    tokens = query.strip().split()  # 简单分词
    # 扩展分词：也在 query 中查找中文词组
    for i in range(len(query)):
        for j in range(i+1, min(i+8, len(query)+1)):
            if len(query[i:j]) >= 2:
                tokens.append(query[i:j])

    # 为每个 token 构建 CONTAINS 条件，用 OR 连接
    conditions = []
    for token in set(tokens):
        conditions.append(f"e.name CONTAINS '{token}'")
    where_clause = " OR ".join(conditions)

    cypher = f"MATCH (e) WHERE {where_clause} RETURN e"

举例说明：
- 用户输入："营销部主管是谁"
- 旧方案（单 CONTAINS）：CONTAINS "营销部主管" → 0 结果（精确词组不存在）
- 新方案（分词 OR）：CONTAINS "营销部" OR CONTAINS "主管" → 先找到营销部节点，再找到主管节点

为什么不用单 CONTAINS？
- 用户输入是自然语言，不是精确的实体名称
- "营销部主管是谁" ≠ 数据库中存的实体名
- 单 CONTAINS 要求子串完全匹配，过于严格
- 分词 + OR 组合：扩大命中范围，再通过 Cypher RETURN 结果让 LLM 筛选解释

安全性保障：
- Cypher 白名单：只允许 MATCH / RETURN / CALL 开头的查询
- 参数化查询：用户输入作为参数传入，防止注入
- 仅读权限：Neo4j 连接使用只读账户""")

# ══════════════════════════════════════════
# 第六章：数字孪生 & 预测 & 告警
# ══════════════════════════════════════════
h1('六、数字孪生 & 预测 & 告警（4 题）')

question('1. 请详细解释"数字孪生"在系统中的完整数据流：从数据录入到 Agent 感知的全过程。', '★★★★☆')
answer("""完整数据流：

Step 1：数据录入
- 方式 A：Settings → Metrics 页面手动单条录入
- 方式 B：POST /api/v1/workspaces/{id}/metrics/batch（JSON 批量导入）

Step 2：数据存储
   INSERT INTO business_metrics (company_id, metric_name, metric_value,
       unit, period, category, recorded_at, tags) VALUES (...)

Step 3：Agent 感知（Orchestrator Step 5.8）
   async with AsyncSessionLocal() as s:
       metrics_svc = BusinessMetricsService(s)
       metrics_ctx = await metrics_svc.get_ai_analysis_context(workspace_id)
       state["context_text"] = metrics_ctx + "\n\n" + existing_context

Step 4：Agent 使用
   CEO Agent 的 System Prompt 收到：
   "Business Metrics Snapshot (company: XXX Corp):
    Monthly Revenue: 1,250,000 元 (period: 2025-06) [revenue]
    Customer Count: 3,420 个 (period: 2025-06) [customer]
    ...
    Total Revenue: 1,250,000.00
    Gross Profit: 470,000.00 (Margin: 37.6%)"

Step 5：Dashboard 展示
   AnalyticsService.get_dashboard_data() 聚合：
   - 指标快照 → KPI 卡片（业务优先级排序）
   - 指标趋势 → 折线图（Recharts）
   - Goals 进度 → 进度条
   - 告警 → 主动告警卡片

Step 6：预测
   用户选择指标 → Forecast 页面 → POST /forecast/predict
   → 读取历史数据 → 移动平均/线性回归 → 返回预测值 + 图表

数据设计要点：
- 每个 metric_name 是多行记录（不同 period），形成时间序列
- category 字段支持分组（revenue/cost/hr/inventory/customer/operations）
- tags JSON 字段支持自定义标签（如 {"region": "华东"}）
- 复合索引 (company_id, metric_name, period) 加速趋势查询""")

# ──
question('2. 预测引擎同时使用移动平均和线性回归，两者各有什么适用场景？为什么自动计算 periods？', '★★★☆☆')
answer("""两种预测方法：

1. 加权移动平均（Weighted Moving Average）：
   - 公式：预测值 = Σ(权重_i × 值_i)
   - 近期数据权重更高（越近越重要）
   - 适用场景：
     a. 数据有季节性波动
     b. 近期趋势比远期更重要
     c. 数据点较少（< 6 个）
   - 优点：平滑噪声，反映近期趋势
   - 缺点：滞后于突变

2. 线性回归（sklearn LinearRegression）：
   - 公式：y = a × t + b
   - 拟合一条最优直线
   - 适用场景：
     a. 数据有明显增长/下降趋势
     b. 数据点较密集
     c. 预测中长期方向
   - 优点：捕捉长期趋势
   - 缺点：对异常值敏感

自动计算 periods 的原因：
   periods = max(1, min(len(values) // 2, 6))
   - min(values//2, 6)：最多预测历史数据量的一半，最多 6 期
   - max(1, ...)：至少预测 1 期
   - 原因：数据点少时不能预测太多期（统计显著性不足）
   - 举例：12 个月历史 → 预测 6 个月；3 个月历史 → 预测 1 个月

去重处理：
   - 每次预测前删除该 metric 的旧预测（避免重复数据点）
   - 预测结果写入 predictions 表，与原始 metrics 分开存储""")

# ──
question('3. Proactive Alerts 的三层检测架构是如何工作的？severity 的评分逻辑是什么？', '★★★★★')
answer("""三层检测架构：

═══ 第 1 层：规则引擎（Rule Engine）═══
- KPI 阈值对比：实际值 vs 目标值
  - 偏差 > 20%：标记为异常
  - 偏差 > 50%：标记为严重异常
- Goals 进度检查：progress_pct < 预期进度 → 告警
- 方向感知：higher-is-better（如收入：达到目标的 80% 正常）
             lower-is-better（如退货率：超过目标的 120% 告警）

═══ 第 2 层：统计异常检测（Statistical Anomaly）═══
- Z-score 方法：|Z| > 2.5 → 统计显著异常
  Z = (当前值 - 均值) / 标准差
- 环比剧变检测：MoM change > 50% → 突增/突降告警
- 趋势反转检测：连续 3 期上涨后突然下跌（或反之）

═══ 第 3 层：AI 评估（LLM Filter）═══
- 将候选告警 + 企业上下文发给 LLM
- LLM 评估："这个异常是否值得告警？"
  - 季节性因素（如 12 月零售额飙升是正常的）
  - 一次性事件（如设备升级导致当月成本突增）
- LLM 返回：{is_alert: true/false, reason: "...", suggestion: "..."}

Severity 评分（4 维加权）：
   score = (
       deviation_score × 40% +  # 偏离程度
       trend_score × 25% +       # 趋势方向
       frequency_score × 15% +   # 历史频率
       ai_score × 20%            # AI 评估
   )
   - critical: score ≥ 70
   - warning: score ≥ 40
   - info: score < 40

告警升级与去重：
- 去重：_is_duplicate() 检查 24 小时内同 metric + 同类型告警
- 升级：info 24h 未读 → warning，warning 72h → critical
- 推送：ws_manager.broadcast_to_workspace() 实时推送告警""")

# ──
question('4. 趋势图中的"上升/下降百分比"是如何计算的？如何处理"越高越好"vs"越低越好"的指标？', '★★★☆☆')
answer("""计算方式（get_trend 端点）：

1. 变化率 = (最新值 - 最早值) / 最早值 × 100%
2. 方向判断：
   change_pct > 0 → "up"
   change_pct < 0 → "down"
   change_pct == 0 → "flat"

"越高越好"指标（如收入、利润率、客户数）：
- 上升 → 绿色箭头 ↑ 正面
- 下降 → 红色箭头 ↓ 负面

"越低越好"指标（如退货率、成本率、缺勤率）：
- Goal.direction = "lower"
- 进度计算反转：progress = target / current × 100（当 current > target）
- 上升 → 红色箭头 ↑ 负面（Alerts 系统标记为告警）
- 下降 → 绿色箭头 ↓ 正面

方向感知实现（_calc_progress）：
   def _calc_progress(current, target, direction):
       if direction == "higher":
           return min(100, current / target * 100)
       else:  # "lower"
           if current <= target: return 100
           return max(0, target / current * 100)

告警检测中的双向感知：
   # 对 lower-is-better 指标也检查比率
   ratio_higher = current / target
   ratio_lower = target / current
   # 取较极端值判断偏离程度
   deviation = max(ratio_higher, ratio_lower)""")

# ══════════════════════════════════════════
# 第七章：部署 & DevOps
# ══════════════════════════════════════════
h1('七、部署 & DevOps（3 题）')

question('1. Docker 多阶段构建在本项目中的应用：前端和后端 Dockerfile 各有什么优化策略？', '★★★★☆')
answer("""前端 Dockerfile（多阶段构建）：

═══ Builder 阶段 ═══
FROM node:22-alpine AS builder
COPY package.json package-lock.json ./
RUN npm ci                          # 仅安装依赖
COPY . .
RUN npm run build                   # Next.js 编译

═══ Runner 阶段 ═══
FROM node:22-alpine AS runner
COPY --from=builder /app/.next/standalone ./  # 仅复制 standalone 产物
COPY --from=builder /app/.next/static ./.next/static

优化策略：
1. 多阶段：编译环境和运行环境分离，最终镜像不包含 node_modules 全部源码
2. Alpine 基础镜像：~50MB vs Debian ~200MB
3. npm ci 而非 npm install：更快、可重现
4. Next.js standalone 模式：自动 tree-shaking，只包含运行时需要的文件
5. 非 root 用户运行（nextjs:nodejs，UID 1001）

后端 Dockerfile 优化策略：
1. CPU-only PyTorch：pip install torch --index-url https://download.pytorch.org/whl/cpu
   避免下载 ~1.5GB 的 CUDA 库（nvidia_cudnn/ nvidia_nccl 等）
2. 阿里云镜像加速：
   - Debian 源：sed 替换 deb.debian.org → mirrors.aliyun.com
   - Pip 源：pip config set global.index-url https://mirrors.aliyun.com/pypi/simple/
3. 最小系统依赖：只安装 curl，去掉 ffmpeg/build-essential
4. Alembic 自动迁移：CMD "alembic upgrade head && uvicorn ..."
   确保容器启动时数据库结构是最新的""")

# ──
question('2. 项目从开发环境迁移到 Docker 部署时遇到了哪些具体问题？是如何解决的？', '★★★★☆')
answer("""实际遇到的问题及解决方案：

1. pip 下载超时（30 分钟超时）：
   原因：sentence-transformers → torch → 10+ NVIDIA CUDA 包（~1.5GB）
   解决：Dockerfile 中先安装 CPU-only PyTorch

2. apt-get 502 Bad Gateway（Debian 源不可达）：
   原因：debian.org 在国内访问不稳定
   解决：sed 替换为 mirrors.aliyun.com

3. .env.local 硬编码 API URL：
   原因：开发时 NEXT_PUBLIC_API_URL=http://localhost:8000，
         打包后浏览器请求仍指向 8000 而非 Nginx 8080
   解决：改为空字符串，使用相对路径

4. passlib + bcrypt 5.0 不兼容：
   原因：bcrypt 5.0 改变了密码长度处理的 API
   解决：添加 bcrypt<5.0 约束

5. celery 未安装导致 backend crash loop：
   原因：requirements.txt 遗漏 celery 依赖
   解决：添加 celery>=5.4.0

6. 端口 80 被 Docker Desktop 自占用：
   原因：Docker Desktop 的 com.docker.backend.exe 监听 80
   解决：Nginx 改为 8080 端口

7. Docker Registry Mirror 不可用：
   原因：docker.1ms.run 返回 403 Forbidden
   解决：Docker Desktop Settings → Docker Engine → 改为 DaoCloud + DockerProxy

8. npm run build lint 错误：
   原因：react/no-unescaped-entities（中文引号、单引号）
   解决：将纯文本引号改为 <b> 标签包裹或移除撇号""")

# ──
question('3. 如果让你给这个系统设计 CI/CD 流水线，你会包含哪些步骤？如何确保部署的安全性？', '★★★☆☆')
answer("""建议的 CI/CD 流水线（GitHub Actions）：

═══ PR 阶段 ═══
1. Lint 检查：
   - 前端：next lint
   - 后端：ruff check / black --check

2. 类型检查：
   - 前端：tsc --noEmit
   - 后端：mypy（逐步引入）

3. 单元测试：
   - 后端：pytest（PostgreSQL testcontainer）
   - 前端：vitest

4. 构建验证：
   - 前端：npm run build（确保编译通过）
   - 后端：alembic check（确保迁移脚本正确）

═══ 合并到 main ═══
5. Docker 镜像构建 + 推送到 Registry
6. 部署到 Staging 环境
7. E2E 测试（Playwright）：
   - 登录 → 创建 Workspace → 录入 Metrics → 发送 Chat → 查看 Analytics
8. 部署到 Production（审批门）

安全性保障：
1. 敏感信息管理：
   - API Key / DB 密码通过 GitHub Secrets 注入
   - Docker 镜像不包含 .env 文件
   - .env.local 加入 .gitignore

2. 依赖扫描：
   - Dependabot 自动 PR 依赖更新
   - pip-audit / npm audit 检查已知漏洞

3. 镜像安全：
   - 使用特定版本标签（不用 latest）
   - 非 root 用户运行容器
   - 最小基础镜像（Alpine/Slim）

4. 网络安全：
   - Nginx 反向代理隐藏后端端口
   - CORS 限制生产环境只允许特定域名
   - HTTPS（生产环境添加 Let's Encrypt）

5. 数据库：
   - 强密码 + 定期轮换
   - 连接池限制（pool_size=20, max_overflow=10）
   - 定期备份（pg_dump cron job）""")

# ══════════════════════════════════════════
# 保存
# ══════════════════════════════════════════
doc.save('EAIOS_Interview_Questions.docx')
print('Saved: EAIOS_Interview_Questions.docx')