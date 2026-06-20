"""生成 EAIOS 完整测试用例 Word 文档（含公司画像 + Metrics 数据集 + 逐步测试步骤）"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.size = Pt(10)
style.font.name = 'Microsoft YaHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

CHAP = dict(size=Pt(28), bold=True)

def h1(text):
    doc.add_page_break()
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.size = Pt(16)
    return h

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.size = Pt(13)
    return h

def h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.size = Pt(11)
    return h

def q(text):
    """步骤标题"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    return p

def _body(t):
    return doc.add_paragraph(t)

def code(code_text, lang=""):
    p = doc.add_paragraph()
    r = p.add_run(code_text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def table(headers, rows, col_widths=None):
    """插入表格"""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return tbl

def tip(text):
    p = doc.add_paragraph()
    r = p.add_run(f"💡 {text}")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    return p

def step(num, title, action, expected=""):
    h3(f"步骤 {num}：{title}")
    q("操作：")
    _body(action)
    if expected:
        q("预期结果：")
        _body(expected)

# ══════════════════════════════════════════
# 封面
# ══════════════════════════════════════════
for _ in range(2): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('EAIOS 完整测试用例手册')
r.font.size = Pt(28); r.bold = True; r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('含公司画像数据 + Metrics 时间序列数据集 + 12 大模块逐步测试步骤')
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph()
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f'v2.0 | {datetime.date.today().strftime("%Y-%m-%d")} | 覆盖 16 页面 × 16 Agent × 8 API 模块').font.size = Pt(10)

doc.add_paragraph()
prep = doc.add_paragraph(); prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
prep.add_run('适用环境：http://localhost:8080 | 测试账号：testadmin / pass123').font.size = Pt(10)

# ══════════════════════════════════════════
# 目录
# ══════════════════════════════════════════
h1('目录')
for item in [
    '第一章   测试准备与环境检查',
    '第二章   企业画像 — 完整测试数据集（可直接复制粘贴）',
    '第三章   Metrics 数据集 — 12 个月时间序列（可直接批量导入）',
    '第四章   测试用例：Settings 模块（5 页）',
    '第五章   测试用例：Document & Knowledge（3 页）',
    '第六章   测试用例：Chat & Multi-Agent（3 场景）',
    '第七章   测试用例：Graph & Memories',
    '第八章   测试用例：Analytics & Forecast',
    '第九章   测试用例：Decision Center',
    '第十章   测试用例：Proactive Alerts',
    '第十一章 测试用例：Reports & Members',
    '第十二章 综合端到端回归测试清单',
]:
    _body(item)

# ══════════════════════════════════════════
# 第一章：准备
# ══════════════════════════════════════════
h1('第一章  测试准备与环境检查')

h2('1.1 启动环境')
code("""# 确保 Docker Desktop 运行中，然后：
cd EAIOS
docker-compose up -d

# 验证 6 个服务全部 healthy
docker ps --format "table {{.Names}}\\t{{.Status}}"
# 预期：workspace-db/redis/neo4j/backend/frontend/nginx 全部 Up""")

h2('1.2 注册测试账号')
code("""打开浏览器访问：http://localhost:8080/login
点击 "Register"
  Email: testadmin@eaios.com
  Username: testadmin
  Password: pass123
点击 "Create Account"
预期：自动跳转至 /chat 页面""")

tip("如果 Username 已被占用，换一个用户名（如 testadmin2）。本节后续步骤假定使用 testadmin 账号。")

h2('1.3 验证基础 API')
code("""curl http://localhost:8080/api/v1/auth/login \\
  -H "Content-Type: application/json" \\
  -d '{"email":"testadmin@eaios.com","password":"pass123"}'
# 预期：返回 {"access_token":"...","refresh_token":"..."} """)

# ══════════════════════════════════════════
# 第二章：企业画像数据
# ══════════════════════════════════════════
h1('第二章  企业画像 — 完整测试数据集')

tip("以下数据用于测试 L1 企业画像的所有 CRUD 功能。按顺序逐模块录入。录入后每个页面刷新可验证数据展示完整性。")

h2('2.1 公司信息（Settings → Company）')
table(
    ['字段', '值', '说明'],
    [
        ['Company Name', '极光科技控股有限公司', '中文全称'],
        ['English Name', 'Aurora Tech Holdings Co., Ltd.', '英文全称'],
        ['Short Name', '极光科技', '简称'],
        ['Industry', '智能制造 / 新能源', '行业'],
        ['Size', '500-1000', '企业规模'],
        ['Founded', '2018', '成立年份'],
        ['HQ Address', '广东省深圳市南山区科技园南路 88 号极光大厦', '总部地址'],
        ['Website', 'https://www.aurora-tech.cn', '官网'],
        ['Description', '极光科技是一家专注于智能光伏逆变器、储能系统和能源管理平台研发制造的高新技术企业。产品远销欧洲、东南亚、中东等 30 余个国家和地区，2024 年营收突破 12 亿元。公司拥有 200+ 研发人员、50+ 核心专利，与清华大学、中科院深度合作。', '企业简介'],
    ]
)

h2('2.2 组织架构（Settings → Organization）')
h3('部门 → 职位 → 员工 三级录入')

q('第一批 — 部门（Departments）')
table(
    ['部门名称', '上级部门', '描述'],
    [
        ['总裁办', '（无，顶级部门）', '战略决策与公司治理'],
        ['研发中心', '总裁办', '产品研发、技术创新、知识产权管理'],
        ['销售与市场部', '总裁办', '全球销售、品牌推广、渠道管理'],
        ['供应链管理部', '总裁办', '采购、生产计划、物流仓储'],
        ['财务与法务部', '总裁办', '财务核算、预算管控、法务合规'],
        ['人力资源部', '总裁办', '招聘、培训、薪酬福利、组织发展'],
        ['客户成功部', '总裁办', '售后技术支持、客户培训、满意度管理'],
        ['软件研发组', '研发中心', '嵌入式软件、云平台、移动 App 开发'],
        ['硬件研发组', '研发中心', '电力电子、结构设计、热管理'],
        ['测试与质量组', '研发中心', '硬件测试、软件测试、可靠性验证'],
        ['国内销售组', '销售与市场部', '国内市场开拓与客户关系'],
        ['海外销售组', '销售与市场部', '海外市场拓展与渠道管理'],
        ['市场品牌组', '销售与市场部', '品牌建设、数字营销、展会策划'],
    ]
)

q('第二批 — 职位（Positions）')
table(
    ['职位名称', '所属部门', '职级'],
    [
        ['首席执行官 CEO', '总裁办', 'L10'],
        ['首席技术官 CTO', '总裁办', 'L9'],
        ['首席财务官 CFO', '总裁办', 'L9'],
        ['研发总监', '研发中心', 'L8'],
        ['销售总监', '销售与市场部', 'L8'],
        ['供应链总监', '供应链管理部', 'L8'],
        ['HR 总监', '人力资源部', 'L8'],
        ['高级软件工程师', '软件研发组', 'L6'],
        ['高级硬件工程师', '硬件研发组', 'L6'],
        ['测试经理', '测试与质量组', 'L7'],
        ['大客户销售经理', '国内销售组', 'L6'],
        ['海外销售经理', '海外销售组', 'L6'],
        ['品牌营销经理', '市场品牌组', 'L5'],
        ['采购经理', '供应链管理部', 'L6'],
        ['客户成功经理', '客户成功部', 'L6'],
        ['财务经理', '财务与法务部', 'L6'],
        ['招聘主管', '人力资源部', 'L5'],
    ]
)

q('第三批 — 员工（Employees）')
table(
    ['姓名', '所属部门', '职位', '邮箱', '电话', '入职日期'],
    [
        ['张明远', '总裁办', '首席执行官 CEO', 'zhangmy@aurora-tech.cn', '13800001001', '2018-03-01'],
        ['王睿之', '总裁办', '首席技术官 CTO', 'wangrz@aurora-tech.cn', '13800001002', '2018-06-15'],
        ['陈婉清', '总裁办', '首席财务官 CFO', 'chenwq@aurora-tech.cn', '13800001003', '2019-01-10'],
        ['李浩然', '研发中心', '研发总监', 'lihr@aurora-tech.cn', '13800002001', '2019-03-01'],
        ['赵雪琳', '销售与市场部', '销售总监', 'zhaoxl@aurora-tech.cn', '13800003001', '2019-06-01'],
        ['刘建国', '供应链管理部', '供应链总监', 'liujg@aurora-tech.cn', '13800004001', '2020-01-15'],
        ['孙丽华', '人力资源部', 'HR 总监', 'sunlh@aurora-tech.cn', '13800005001', '2020-03-01'],
        ['周文博', '软件研发组', '高级软件工程师', 'zhouwb@aurora-tech.cn', '13800002002', '2021-07-01'],
        ['吴振宇', '硬件研发组', '高级硬件工程师', 'wuzhy@aurora-tech.cn', '13800002003', '2021-09-01'],
        ['郑晓燕', '测试与质量组', '测试经理', 'zhengxy@aurora-tech.cn', '13800002004', '2021-11-15'],
        ['黄志强', '国内销售组', '大客户销售经理', 'huangzq@aurora-tech.cn', '13800003002', '2022-01-01'],
        ['林雨桐', '海外销售组', '海外销售经理', 'linyt@aurora-tech.cn', '13800003003', '2022-03-01'],
        ['何美玲', '市场品牌组', '品牌营销经理', 'heml@aurora-tech.cn', '13800003004', '2022-05-01'],
        ['马伟东', '供应链管理部', '采购经理', 'mawd@aurora-tech.cn', '13800004002', '2022-07-01'],
        ['徐静怡', '客户成功部', '客户成功经理', 'xujy@aurora-tech.cn', '13800006001', '2023-01-01'],
        ['高志远', '财务与法务部', '财务经理', 'gaozy@aurora-tech.cn', '13800007001', '2023-03-01'],
        ['林晓燕', '人力资源部', '招聘主管', 'linxy@aurora-tech.cn', '13800005002', '2023-06-01'],
    ]
)

h2('2.3 业务信息（Settings → Business）')
h3('产品（Products）')
table(
    ['产品名称', '类型', '价格(元)', '描述'],
    [
        ['Aurora S5000 光伏逆变器', '硬件', '35000', '5kW 单相家用光伏逆变器，效率 98.5%，IP65 防护，WiFi/4G 双模通信'],
        ['Aurora S10000 光伏逆变器', '硬件', '65000', '10kW 三相商用逆变器，最大效率 99.0%，支持 AFCI 电弧检测'],
        ['Aurora B10 储能电池', '硬件', '28000', '10kWh 家用储能锂电池组，循环寿命 6000 次，10 年质保'],
        ['Aurora Cloud EMS', '软件', '9999/年', '云端能源管理系统，实时监控、智能调度、AI 预测维护'],
        ['Aurora Mobile App', '软件', '免费', '用户端 App，发电量查看、储能调度、故障告警推送'],
        ['Aurora Commercial EMS', '软件', '49999/年', '商用能源管理系统，多站点管理、负荷预测、VPP 聚合'],
    ]
)

h3('客户（Customers）')
table(
    ['客户名称', '行业', '所在地区', '客户等级', '联系人', '年采购额(万元)'],
    [
        ['阳光新能源有限公司', '光伏电站', '华东', 'A 级', '刘总', '1500'],
        ['绿能时代科技有限公司', '储能系统', '华南', 'A 级', '陈总', '1200'],
        ['德盛电力工程有限公司', '电力工程', '华北', 'B 级', '王总', '800'],
        ['EuroGreen Energy GmbH', '新能源分销', '德国', 'A 级', 'Dr. Mueller', '2000'],
        ['SolaireTech SAS', '光伏安装', '法国', 'B 级', 'M. Dupont', '600'],
        ['中东新能源集团', '综合能源', '阿联酋', 'A 级', 'Mr. Al-Rashid', '1800'],
    ]
)

h3('业务流程（Processes）')
table(
    ['流程名称', '所属部门', '步骤数', '描述'],
    [
        ['新产品开发流程 NPD', '研发中心', '7', '市场调研→概念设计→详细设计→原型验证→小批量试产→量产评审→上市'],
        ['销售订单处理流程', '销售与市场部', '5', '客户询价→报价审批→合同签订→订单录入→交付跟踪'],
        ['供应商评估流程', '供应链管理部', '4', '供应商筛选→样品验证→现场审核→年度评估'],
        ['客户投诉处理流程', '客户成功部', '5', '投诉接收→问题分类→技术分析→解决方案→客户回访闭环'],
        ['员工招聘流程', '人力资源部', '6', '需求确认→JD发布→简历筛选→多轮面试→Offer→入职'],
    ]
)

h2('2.4 Goals & KPIs（Settings → Goals）')
table(
    ['目标名称', '类型', '目标值', '当前值', '方向', '时间范围', '描述'],
    [
        ['年度营业收入', 'Revenue', '1500000000', '750000000', 'higher', '2025-01~2025-12', '全年营收目标 15 亿（元），半年度进度'],
        ['产品毛利率', 'Profitability', '42', '39.5', 'higher', '2025-Q2', '目标毛利率 42%（%）'],
        ['客户满意度 NPS', 'Customer', '65', '58', 'higher', '2025-Q2', 'NPS 净值目标 65'],
        ['产品退货率', 'Quality', '1.5', '1.8', 'lower', '2025-Q2', '控制退货率在 1.5% 以下'],
        ['员工流失率', 'HR', '10', '7.2', 'lower', '2025-Q2', '年度员工主动离职率 < 10%'],
        ['新增海外客户数', 'Growth', '30', '16', 'higher', '2025-01~2025-12', '全年新增海外签约客户 ≥ 30 家'],
        ['研发投入占比', 'Investment', '12', '10.8', 'higher', '2025-Q2', '研发投入占营收 12%'],
    ]
)

# ══════════════════════════════════════════
# 第三章：Metrics 数据集
# ══════════════════════════════════════════
h1('第三章  Metrics 数据集 — 12 个月时间序列')

tip("这是数字孪生的核心数据。以下 JSON 可直接在 Settings → Metrics 页面点击 'Batch Import' 粘贴导入。")

h2('3.1 批量导入 JSON（一键复制）')

json_data = """[
  {"metric_name":"Monthly Revenue","metric_value":82000000,"unit":"元","period":"2024-07","category":"revenue","notes":"Q3开局稳定"},
  {"metric_name":"Monthly Revenue","metric_value":85000000,"unit":"元","period":"2024-08","category":"revenue","notes":"暑期旺季增长"},
  {"metric_name":"Monthly Revenue","metric_value":88000000,"unit":"元","period":"2024-09","category":"revenue","notes":"Q3冲量"},
  {"metric_name":"Monthly Revenue","metric_value":92000000,"unit":"元","period":"2024-10","category":"revenue","notes":"Q4开局"},
  {"metric_name":"Monthly Revenue","metric_value":105000000,"unit":"元","period":"2024-11","category":"revenue","notes":"双十一促销"},
  {"metric_name":"Monthly Revenue","metric_value":115000000,"unit":"元","period":"2024-12","category":"revenue","notes":"年末旺季"},
  {"metric_name":"Monthly Revenue","metric_value":95000000,"unit":"元","period":"2025-01","category":"revenue","notes":"春节淡季"},
  {"metric_name":"Monthly Revenue","metric_value":98000000,"unit":"元","period":"2025-02","category":"revenue","notes":"节后恢复"},
  {"metric_name":"Monthly Revenue","metric_value":102000000,"unit":"元","period":"2025-03","category":"revenue","notes":"Q1冲量"},
  {"metric_name":"Monthly Revenue","metric_value":108000000,"unit":"元","period":"2025-04","category":"revenue","notes":""},
  {"metric_name":"Monthly Revenue","metric_value":112000000,"unit":"元","period":"2025-05","category":"revenue","notes":""},
  {"metric_name":"Monthly Revenue","metric_value":125000000,"unit":"元","period":"2025-06","category":"revenue","notes":"H1最高"},
  {"metric_name":"Monthly Cost","metric_value":52000000,"unit":"元","period":"2024-07","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":53000000,"unit":"元","period":"2024-08","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":54000000,"unit":"元","period":"2024-09","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":56000000,"unit":"元","period":"2024-10","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":62000000,"unit":"元","period":"2024-11","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":68000000,"unit":"元","period":"2024-12","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":58000000,"unit":"元","period":"2025-01","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":60000000,"unit":"元","period":"2025-02","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":62000000,"unit":"元","period":"2025-03","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":65000000,"unit":"元","period":"2025-04","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":68000000,"unit":"元","period":"2025-05","category":"cost"},
  {"metric_name":"Monthly Cost","metric_value":75000000,"unit":"元","period":"2025-06","category":"cost"},
  {"metric_name":"Customer Count","metric_value":2850,"unit":"个","period":"2024-07","category":"customer"},
  {"metric_name":"Customer Count","metric_value":2900,"unit":"个","period":"2024-08","category":"customer"},
  {"metric_name":"Customer Count","metric_value":2980,"unit":"个","period":"2024-09","category":"customer"},
  {"metric_name":"Customer Count","metric_value":3050,"unit":"个","period":"2024-10","category":"customer"},
  {"metric_name":"Customer Count","metric_value":3120,"unit":"个","period":"2024-11","category":"customer"},
  {"metric_name":"Customer Count","metric_value":3200,"unit":"个","period":"2024-12","category":"customer"},
  {"metric_name":"Customer Count","metric_value":3150,"unit":"个","period":"2025-01","category":"customer"},
  {"metric_name":"Customer Count","metric_value":3180,"unit":"个","period":"2025-02","category":"customer"},
  {"metric_name":"Customer Count","metric_value":3250,"unit":"个","period":"2025-03","category":"customer"},
  {"metric_name":"Customer Count","metric_value":3320,"unit":"个","period":"2025-04","category":"customer"},
  {"metric_name":"Customer Count","metric_value":3380,"unit":"个","period":"2025-05","category":"customer"},
  {"metric_name":"Customer Count","metric_value":3420,"unit":"个","period":"2025-06","category":"customer"},
  {"metric_name":"Employee Count","metric_value":680,"unit":"人","period":"2024-07","category":"hr"},
  {"metric_name":"Employee Count","metric_value":690,"unit":"人","period":"2024-10","category":"hr"},
  {"metric_name":"Employee Count","metric_value":720,"unit":"人","period":"2025-01","category":"hr"},
  {"metric_name":"Employee Count","metric_value":750,"unit":"人","period":"2025-04","category":"hr"},
  {"metric_name":"R&D Headcount","metric_value":180,"unit":"人","period":"2024-07","category":"hr"},
  {"metric_name":"R&D Headcount","metric_value":195,"unit":"人","period":"2024-10","category":"hr"},
  {"metric_name":"R&D Headcount","metric_value":210,"unit":"人","period":"2025-01","category":"hr"},
  {"metric_name":"R&D Headcount","metric_value":225,"unit":"人","period":"2025-04","category":"hr"},
  {"metric_name":"Inventory Value","metric_value":45000000,"unit":"元","period":"2024-07","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":48000000,"unit":"元","period":"2024-08","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":52000000,"unit":"元","period":"2024-09","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":49000000,"unit":"元","period":"2024-10","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":55000000,"unit":"元","period":"2024-11","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":42000000,"unit":"元","period":"2024-12","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":38000000,"unit":"元","period":"2025-01","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":41000000,"unit":"元","period":"2025-02","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":44000000,"unit":"元","period":"2025-03","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":47000000,"unit":"元","period":"2025-04","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":50000000,"unit":"元","period":"2025-05","category":"inventory"},
  {"metric_name":"Inventory Value","metric_value":48000000,"unit":"元","period":"2025-06","category":"inventory"},
  {"metric_name":"Order Fulfillment Rate","metric_value":94.5,"unit":"%","period":"2024-07","category":"operations"},
  {"metric_name":"Order Fulfillment Rate","metric_value":95.0,"unit":"%","period":"2024-10","category":"operations"},
  {"metric_name":"Order Fulfillment Rate","metric_value":96.2,"unit":"%","period":"2025-01","category":"operations"},
  {"metric_name":"Order Fulfillment Rate","metric_value":96.8,"unit":"%","period":"2025-04","category":"operations"},
  {"metric_name":"Production Capacity Utilization","metric_value":78,"unit":"%","period":"2024-07","category":"operations"},
  {"metric_name":"Production Capacity Utilization","metric_value":82,"unit":"%","period":"2024-10","category":"operations"},
  {"metric_name":"Production Capacity Utilization","metric_value":85,"unit":"%","period":"2025-01","category":"operations"},
  {"metric_name":"Production Capacity Utilization","metric_value":88,"unit":"%","period":"2025-04","category":"operations"}
]"""

code(json_data)

tip("共 60 条 Metrics 数据，覆盖 8 个指标 × 6 个分类（revenue/cost/customer/hr/inventory/operations），跨度 12 个月。导入后可直接用于测试 Analytics / Forecast / Digital Twin 功能。")

h2('3.2 预期分析结果')
table(
    ['指标', '最新值 (2025-06)', '12 月变化', '趋势'],
    [
        ['Monthly Revenue', '1.25 亿元', '+52.4% (从 8200 万)', '↑ 持续上升'],
        ['Monthly Cost', '7500 万元', '+44.2% (从 5200 万)', '↑ 上升中（需关注成本增速）'],
        ['Customer Count', '3420 个', '+20.0% (从 2850)', '↑ 稳定增长'],
        ['Employee Count', '750 人', '+10.3% (从 680)', '↑ 稳步扩张'],
        ['Inventory Value', '4800 万元', '+6.7% (从 4500 万)', '→ 小幅波动'],
        ['Order Fulfillment Rate', '96.8%', '+2.3pp (从 94.5%)', '↑ 持续改善'],
    ]
)

# ══════════════════════════════════════════
# 第四章：Settings 模块测试
# ══════════════════════════════════════════
h1('第四章  测试用例：Settings 模块')
tip("Settings 模块共 5 个子页面：Company / Organization / Business / Goals / Metrics。以下逐一测试。")

# ── Company ──
h2('4.1 Settings → Company')
step(1, '创建企业信息',
    '访问 /settings/company → 按第二章 2.1 的表格填入所有字段 → 点击 Save',
    '页面显示所有已填字段的预览（非编辑态）；Company Name 出现在顶部')
step(2, 'Edit 修改企业信息',
    '点击 Edit 按钮 → 修改 Description 字段 → 点击 Save → 刷新页面',
    'Description 更新为新值，其他字段不变')
step(3, '空字段验证',
    '点击 Edit → 清空 Company Name → 尝试保存',
    '前端提示"必填字段"或后端返回 400 错误')

# ── Organization ──
h2('4.2 Settings → Organization')
step(1, '创建部门（含层级）',
    '按第二章 2.2 表格逐行录入 Department → Name + Parent + Description',
    '树形结构正确展示：总裁办 → 研发中心/销售市场部/... → 子部门')
step(2, '创建职位并关联部门',
    '切换到 Positions Tab → 逐行录入第二章 2.2 的职位数据',
    '每个职位显示其所属部门名')
step(3, '创建员工并关联部门+职位',
    '切换到 Employees Tab → 逐行录入第二章 2.2 的员工数据',
    '每个员工显示所属部门和职位')
step(4, 'Edit 员工信息',
    '点击某个员工的 Edit 图标 → 修改电话号码 → 保存',
    '员工信息更新，部门和职位关联不变')
step(5, 'Delete 员工',
    '点击某个员工的 Delete 图标 → 确认删除',
    '员工从列表中消失，其他员工不受影响')

# ── Business ──
h2('4.3 Settings → Business')
step(1, '创建产品',
    'Products Tab → 录入第二章 2.3 的 6 条产品数据',
    '产品列表中正确展示名称/价格/描述')
step(2, '创建客户',
    'Customers Tab → 录入第二章 2.3 的 6 条客户数据',
    '客户列表展示名称/地区/等级/年采购额')
step(3, '创建业务流程',
    'Processes Tab → 录入第二章 2.3 的 5 条流程数据',
    '流程显示所属部门和步骤数')
step(4, 'Edit 修改',
    '分别对 Products/Customers/Processes 点击 Edit → 修改 → 保存',
    '每次修改后数据正确更新')
step(5, '信息预览完整性',
    '在非编辑态下浏览所有条目',
    '每种记录的关键字段都完整显示（不截断）')

# ── Goals ──
h2('4.4 Settings → Goals')
step(1, '创建 Goals',
    '录入第二章 2.4 的 7 条 Goal → 注意正确选择 direction（higher/lower）',
    '列表正确展示 Goal 名称/目标值/当前值/direction')
step(2, 'Edit Goal',
    '修改"年度营业收入"的 current_value 从 750000000 → 800000000 → 保存',
    'Current Value 更新，进度条重新计算')
step(3, 'Lower-is-better 的 progress 计算验证',
    '查看"产品退货率"（target=1.5, current=1.8, direction=lower）的 progress',
    'progress = 1.5/1.8×100 ≈ 83%（正确使用 lower 方向公式）')

# ── Metrics ──
h2('4.5 Settings → Metrics')
step(1, '批量导入 60 条 Metrics',
    '点击 Batch Import → 粘贴第三章的完整 JSON → 点击 Save',
    '成功导入 60 条；页面显示最新录入的指标列表')
step(2, '验证列表排序',
    '观察 Metrics 列表',
    '排序规则：category ASC → metric_name ASC → recorded_at DESC')
step(3, 'Edit 单条 Metric',
    '找到任意一条 → 点击 Edit → 修改 value 和 notes → 保存',
    '数据更新')
step(4, '逐月补录',
    '使用 Single Record 表单手动追加一条 2025-07 的 Monthly Revenue = 130000000',
    '成功后 metric 列表多一条')

# ══════════════════════════════════════════
# 第五章：Document & Knowledge
# ══════════════════════════════════════════
h1('第五章  测试用例：Documents & Knowledge')

h2('5.1 Documents 页')
step(1, '上传文档',
    '访问 /documents → 点击 Upload → 选择一个 PDF/Word/PPTX/TXT 文件',
    '文件出现在文档列表中，状态为"已处理"或"处理中"')
step(2, '多格式测试',
    '分别上传 PDF、DOCX、PPTX、XLSX、TXT、MD 格式各 1 个',
    '6 个文档全部成功解析入库')
step(3, '文档预览',
    '点击任意文档 → 查看预览',
    '显示文档内容摘要或完整文本')
step(4, '删除文档',
    '点击文档的 Delete 按钮 → 确认',
    '文档从列表中移除')

h2('5.2 Knowledge 页')
step(1, '语义搜索',
    '访问 /knowledge → 在搜索框输入"光伏逆变器的技术参数" → 搜索',
    '返回相关文档片段 + 来源引用（文档名 + Chunk 位置）')
step(2, '跨文档检索',
    '搜索"成本控制" → 观察结果',
    '结果可能来自多个文档（如财务报告 + 会议纪要）')
step(3, '连接器 Tab 验证',
    '分别切换到 GitHub/Notion/Jira/Confluence 各 Tab',
    '每个 Tab 显示对应的配置界面（API Key 输入 + 连接按钮）')

# ══════════════════════════════════════════
# 第六章：Chat & Multi-Agent
# ══════════════════════════════════════════
h1('第六章  测试用例：Chat & Multi-Agent')

h2('6.1 基础对话流')
step(1, '普通对话',
    '访问 /chat → 输入"你好，介绍一下你自己" → 发送',
    'Agent 回复，介绍自己是企业 AI 助手')
step(2, '流式输出',
    '输入一个需要较长回答的问题 → 观察输出方式',
    '逐字/逐 token 流式渲染，非一次性弹出')
step(3, '多轮对话',
    '连续输入 5 条相关消息',
    'Agent 保持上下文一致，能引用之前的对话')

h2('6.2 Multi-Agent 路由测试')
step(1, '企业画像 Agent（profile）',
    '输入："我们公司有哪些部门？"',
    'Supervisor 路由到 profile Agent，回答列出你录入的 13 个部门')
step(2, '财务 Agent（finance）→ 需先录入 Metrics',
    '输入："分析我们公司最近半年的收入趋势"',
    'Supervisor 路由到 finance Agent，引用你录入的 Monthly Revenue 数据')
step(3, '销售 Agent（sales）',
    '输入："我们目前有多少客户？海外市场表现如何？"',
    'Supervisor 路由到 sales Agent，引用 Customer Count 和客户数据')
step(4, 'HR Agent',
    '输入："公司员工增长情况如何？研发人员占比多少？"',
    'Supervisor 路由到 hr Agent，引用 Employee Count & R&D Headcount')
step(5, 'CEO Agent（战略）',
    '输入："综合评估公司当前经营状况，给出改进建议"',
    'CEO Agent 综合收入/成本/客户/员工等多维度数据给出战略建议')
step(6, '多 Agent 流水线',
    '输入："搜索文档中的财务数据，做分析，然后生成一份报告"',
    'Agent Trace 显示 search → analyst → writer 顺序执行')

h2('6.3 上下文感知验证')
step(1, '企业画像上下文',
    '在上传企业信息后 → 问 CEO Agent："我们公司的核心产品是什么？"',
    'Agent 引用 Aurora 系列产品信息')
step(2, 'Metrics 上下文',
    '在录入 Metrics 后 → 问 finance Agent："我们上个月收入多少？"',
    'Agent 回答 2025-06 的 Monthly Revenue = 1.25 亿')
step(3, 'Memory 上下文（需先有对话历史）',
    '经过多轮对话后 → 问 memory Agent："我们之前讨论过什么重要决定？"',
    'Agent 召回之前的对话记忆')

# ══════════════════════════════════════════
# 第七章：Graph & Memories
# ══════════════════════════════════════════
h1('第七章  测试用例：Graph & Memories')

h2('7.1 Knowledge Graph')
step(1, '图谱同步',
    '访问 /graph → 点击 Sync（同步）按钮',
    'Neo4j 从企业画像数据中创建实体和关系节点')
step(2, '图谱搜索',
    '搜索："张明远" → 查看结果',
    '返回张明远节点：CEO / 总裁办 / 关联的部门')
step(3, '关系查询',
    '搜索："谁在研发中心" → 查看结果',
    '返回研发中心及其下属部门的所有人')
step(4, '分词搜索验证',
    '搜索："营销部主管"',
    '分词为["营销部", "主管"] → OR 匹配 → 即使没有精确"营销部主管"实体也能找到相关结果')
step(5, '图谱统计',
    '点击 Stats（统计）',
    '显示实体总数 / 关系总数 / 节点类型分布')

h2('7.2 Memories 页')
step(1, '记忆列表',
    '访问 /memories → 查看记忆列表',
    '显示 Long-Term / Episodic / Semantic 三类记忆')
step(2, 'Event Timeline',
    '查看时间线',
    '事件按时间戳降序排列（最新在前）')
step(3, '记忆召回',
    '输入关键词搜索',
    '混合加权召回（importance x 40% + recency x 30% + vector x 30%）返回相关记忆')

# ══════════════════════════════════════════
# 第八章：Analytics & Forecast
# ══════════════════════════════════════════
h1('第八章  测试用例：Analytics & Forecast')

h2('8.1 Analytics Dashboard')
step(1, '仪表盘加载',
    '访问 /analytics → 等待数据加载',
    '页面展示 KPI 卡片 + 趋势图 + Goals 进度条 + 告警面板')
step(2, 'KPI 卡片排序',
    '观察 KPI 卡片排列顺序',
    'revenue → cost → hr → inventory → customer → operations 的业务优先级排序')
step(3, '趋势图渲染',
    '检查每个指标的折线图',
    '12 个数据点正确绘制；悬停显示 tooltip；Target 虚线正确显示')
step(4, 'AI 分析面板',
    '滚动到 AI Analysis 区域',
    'LLM 生成摘要 + 关键发现 + 建议；缓存 1 小时有效')
step(5, 'Proactive Alert 集成',
    '查看 Alert 区域',
    '如果有告警，显示 Bell 图标 + 告警数量；点击展开详情')
step(6, 'Check Now 按钮',
    '点击 Check Now 按钮 → 等待结果',
    '显示当前检测状态；新告警推送到面板')

h2('8.2 Forecast 页')
step(1, '选择指标并预测',
    '访问 /forecast → 从下拉选择"Monthly Revenue" → 选择方法"Linear Regression" → 点击 Predict',
    '展示 12 个历史数据点 + N 个预测数据点；N = min(历史数/2, 6) = 6')
step(2, '切换预测方法',
    '切换到"Moving Average" → 重新 Predict',
    '两种方法给出不同的预测结果；MA 更平滑，LR 更趋势化')
step(3, '多指标预测',
    '对"Monthly Cost"、"Customer Count"也分别预测',
    '各指标的预测结果独立展示')
step(4, '去重验证',
    '对同一指标连续 Predict 两次',
    '不会出现重复的预测数据点')

# ══════════════════════════════════════════
# 第九章：Decision Center
# ══════════════════════════════════════════
h1('第九章  测试用例：Decision Center')

step(1, '发起战略分析',
    '访问 /decision-center → 输入："公司是否应该进入东南亚储能市场？" → 提交',
    'SSE 流式返回分析结果，Markdown 渲染；5 维度框架覆盖')
step(2, 'Decision History',
    '点击 Decision History 展开',
    '显示之前的所有决策分析记录')
step(3, '展开/折叠历史',
    '点击某条历史记录 → 展开 → 再次点击折叠',
    'Markdown 内容正确渲染')
step(4, '删除决策记录',
    '点击某条记录的 Delete 按钮（需 Admin 权限）',
    '记录被删除，列表更新')
step(5, 'Decision Center 不污染 Chat',
    '访问 /chat → 查看对话列表',
    'Decision Center 的分析请求不出现在 Chat 历史中')

# ══════════════════════════════════════════
# 第十章：Proactive Alerts
# ══════════════════════════════════════════
h1('第十章  测试用例：Proactive Alerts')

step(1, '触发指标告警',
    '手动修改某条 KPI 对应的 Metric 值使其严重偏离目标（如 Monthly Revenue 改为 50000000）→ 等待后台 Celery Beat 检测（或手动调用 Check Now）',
    '系统生成 Proactive Alert')
step(2, '告警展示',
    '查看 Analytics 页面的 Alert 面板',
    '告警卡片：severity 颜色 + 类型图标 + 描述 + 建议')
step(3, '告警去重',
    '再次触发同类型的 Check → 观察 Alert 面板',
    '24 小时内同指标同类型不重复生成')
step(4, '告警 Dismiss',
    '点击某个 Alert 的 Dismiss 按钮',
    'Alert 标记为已读/已忽略，不再显示')
step(5, '告警升级',
    '等待 info 告警超过 24 小时未读（或手动修改时间）',
    '自动升级为 warning → critical')
step(6, 'WebSocket 推送',
    '保持 Analytics 页面打开 → 触发新告警',
    '页面实时收到推送通知（无需刷新）')

# ══════════════════════════════════════════
# 第十一章：Reports & Members
# ══════════════════════════════════════════
h1('第十一章  测试用例：Reports & Members')

h2('11.1 Reports 页')
step(1, '生成报告',
    '访问 /reports → 点击 Generate Report',
    '系统后台生成报告')
step(2, '报告列表',
    '查看报告列表',
    '显示已生成的报告（名称、日期、状态）')
step(3, '报告下载',
    '点击某份报告的 Download 按钮',
    '下载文件为 PDF/DOCX 格式')

h2('11.2 Members 页')
step(1, '成员列表',
    '访问 /members',
    '显示当前 Workspace 的所有成员 + 角色')
step(2, '搜索成员',
    '在搜索框输入"admin" → 结果过滤',
    '只显示匹配的成员')
step(3, '添加成员',
    '输入邮箱 → 选择角色（Admin/Member/Viewer） → 点击 Add',
    '新成员添加成功 → 该用户获得对应权限的 Workspace 访问')
step(4, '权限验证',
    '用 Viewer 角色的账号登录 → 尝试删除 Metrics',
    '返回 403 Forbidden（Viewer 不能删除）')

# ══════════════════════════════════════════
# 第十二章：综合回归测试清单
# ══════════════════════════════════════════
h1('第十二章  综合端到端回归测试清单')

h2('12.1 完整用户旅程（Smoke Test）')
steps_smoke = [
    '1. 注册新用户 → 登录 → 跳转到 /chat',
    '2. 创建 Company Profile → Settings → Company（填入完整信息）',
    '3. 录入 Organization：3 级部门 → 职位 → 员工（17 条）',
    '4. 录入 Business：6 Products + 6 Customers + 5 Processes',
    '5. 设定 Goals & KPIs（7 条，含 higher/lower 两种方向）',
    '6. 批量导入 Metrics（60 条，12 个月 × 8 指标）',
    '7. 上传 3 个不同格式的文档',
    '8. 在 Chat 中发送 10 条消息，涵盖 5 种 Agent（profile/finance/sales/hr/ceo）',
    '9. 检查 Analytics Dashboard 数据完整性',
    '10. 在 Forecast 页面预测 3 个指标',
    '11. 在 Graph 页面同步并搜索',
    '12. 在 Decision Center 发起 1 次战略分析',
    '13. 检查 Alerts 面板',
    '14. 在 Members 页面添加 1 个 Viewer → 验证权限限制',
]
for s in steps_smoke:
    _body(s)

h2('12.2 API 可用性检查')
code("""# 1. Health Check
curl http://localhost:8080/api/v1/

# 2. Auth
curl -X POST http://localhost:8080/api/v1/auth/login -H "Content-Type: application/json" -d '{"email":"testadmin@eaios.com","password":"pass123"}'

# 3. Metrics Snapshot
curl -H "Authorization: Bearer $TOKEN" http://localhost:8080/api/v1/workspaces/{id}/metrics/snapshot

# 4. Analytics Dashboard
curl -H "Authorization: Bearer $TOKEN" http://localhost:8080/api/v1/workspaces/{id}/analytics/dashboard

# 5. Graph Query
curl -H "Authorization: Bearer $TOKEN" "http://localhost:8080/api/v1/workspaces/{id}/graph/search?q=研发中心"

# 6. Forecast
curl -X POST -H "Authorization: Bearer $TOKEN" -H "Content-Type: application/json" \\
  http://localhost:8080/api/v1/workspaces/{id}/forecast/predict \\
  -d '{"metric_name":"Monthly Revenue","method":"linear_regression"}'""")

h2('12.3 前端页面可访问性')
pages = [
    '/chat', '/documents', '/knowledge', '/graph', '/memories',
    '/analytics', '/forecast', '/decision-center', '/reports', '/members',
    '/settings/company', '/settings/org', '/settings/business', '/settings/goals', '/settings/metrics',
    '/login'
]
for page in pages:
    _body(f"  □ {page}")

h2('12.4 非功能验证')
table(
    ['项目', '验证方法', '期望'],
    [
        ['响应式布局', '浏览器窗口缩至 < 768px', 'Hamburger 菜单出现；侧边栏 overlay 显示'],
        ['404 页面', '访问 /nonexistent-page', '自定义 404 页面（品牌化设计 + 返回首页按钮）'],
        ['500 错误处理', '手动触发后端异常', 'error.tsx 展示错误信息 + 重试按钮'],
        ['Token 自动刷新', '等待 60 分钟后操作', 'Token 自动刷新，用户无需重新登录'],
        ['移动端响应式', 'Chrome DevTools 模拟 iPhone 14', '所有页面可正常操作'],
    ]
)

# ══════════════════════════════════════════
# 保存
# ══════════════════════════════════════════
doc.save('EAIOS_Test_Cases.docx')
print('Saved: EAIOS_Test_Cases.docx')